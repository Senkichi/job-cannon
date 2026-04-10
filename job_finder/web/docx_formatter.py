"""Resume .docx formatter using python-docx.

Produces a professionally formatted .docx from a structured resume dict.
Uses only standard Word styles (Heading 1, Heading 2, List Bullet, Normal)
to ensure reliable conversion to Google Docs format.

Research note (Pitfall 5): Custom styles are not preserved during Google Docs
conversion. Stick to built-in Word styles for best fidelity.

Usage:
    from job_finder.web.docx_formatter import build_resume_docx

    buffer = build_resume_docx(resume_data)
    url = upload_to_drive(service, "Resume - Acme Corp", buffer, folder_id="...")

Expected resume_data structure:
    {
        "name": str,
        "contact_line": str,           # "email | phone | linkedin | location"
        "summary": str,
        "skills": list[str],
        "positions": list[{
            "title": str,
            "company": str,
            "dates": str,
            "achievements": list[str],
        }],
        "education": list[{
            "degree": str,
            "institution": str,
            "year": str,
        }],
    }
"""

import io
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


# ATS-safe character normalization map.
# ATS parsers (Workday, Taleo, iCIMS) choke on these Unicode characters
# during keyword extraction. Replace with ASCII equivalents.
_ATS_NORMALIZE_MAP = str.maketrans({
    "\u2018": "'",     # left single quote
    "\u2019": "'",     # right single quote (apostrophe)
    "\u201A": "'",     # single low-9 quote
    "\u201B": "'",     # single high-reversed-9 quote
    "\u201C": '"',     # left double quote
    "\u201D": '"',     # right double quote
    "\u201E": '"',     # double low-9 quote
    "\u201F": '"',     # double high-reversed-9 quote
    "\u2014": " - ",   # em dash -> space-hyphen-space
    "\u2013": "-",     # en dash -> hyphen
    "\u2026": "...",   # ellipsis
    "\u00A0": " ",     # non-breaking space
    "\u200B": "",      # zero-width space (remove)
    "\u200C": "",      # zero-width non-joiner (remove)
    "\u200D": "",      # zero-width joiner (remove)
    "\uFEFF": "",      # BOM / zero-width no-break space (remove)
    "\u2022": "-",     # bullet -> hyphen (for inline lists)
    "\u25CF": "-",     # black circle bullet
    "\u25CB": "-",     # white circle bullet
    "\u00B7": "-",     # middle dot
    "\u2023": "-",     # triangular bullet
    "\u00AB": '"',     # left guillemet
    "\u00BB": '"',     # right guillemet
    "\u2039": "'",     # single left angle quote
    "\u203A": "'",     # single right angle quote
})


def _normalize_for_ats(text: str) -> str:
    """Replace Unicode characters that break ATS keyword matching."""
    if not text:
        return text
    return text.translate(_ATS_NORMALIZE_MAP)


def _normalize_resume_data(data):
    """Recursively normalize all strings in resume_data for ATS compatibility."""
    if isinstance(data, str):
        return _normalize_for_ats(data)
    if isinstance(data, list):
        return [_normalize_resume_data(item) for item in data]
    if isinstance(data, dict):
        return {k: _normalize_resume_data(v) for k, v in data.items()}
    return data


def _set_margins(doc: Document, margin_inches: float = 1.0) -> None:
    """Set uniform page margins on all sections."""
    for section in doc.sections:
        section.top_margin = Inches(margin_inches)
        section.bottom_margin = Inches(margin_inches)
        section.left_margin = Inches(margin_inches)
        section.right_margin = Inches(margin_inches)

def _add_centered_paragraph(doc: Document, text: str, style: str, font_size: Optional[int] = None) -> None:
    """Add a centered paragraph with the given style and optional font size."""
    para = doc.add_paragraph(text, style=style)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if font_size is not None:
        for run in para.runs:
            run.font.size = Pt(font_size)

def build_resume_docx(resume_data: dict) -> io.BytesIO:
    """Build a .docx resume document from structured data.

    Creates a professionally formatted Word document with sections for
    name/contact, professional summary, technical skills, experience, and
    education. All styles are standard Word built-in styles for reliable
    Google Docs conversion.

    Args:
        resume_data: Structured resume dict. See module docstring for schema.

    Returns:
        BytesIO positioned at 0 containing the .docx file bytes.
    """
    resume_data = _normalize_resume_data(resume_data)

    doc = Document()
    _set_margins(doc, margin_inches=1.0)

    # --- Name (Heading 1, centered) ---
    _add_centered_paragraph(doc, resume_data.get("name", ""), style="Heading 1")

    # --- Contact line (Normal, centered, 10pt) ---
    contact_line = resume_data.get("contact_line", "")
    if contact_line:
        _add_centered_paragraph(doc, contact_line, style="Normal", font_size=10)

    # --- Professional Summary ---
    doc.add_paragraph("Professional Summary", style="Heading 2")
    summary = resume_data.get("summary", "")
    if summary:
        doc.add_paragraph(summary, style="Normal")

    # --- Technical Skills ---
    doc.add_paragraph("Technical Skills", style="Heading 2")
    skills = resume_data.get("skills", [])
    if skills:
        para = doc.add_paragraph(style="Normal")
        bold_run = para.add_run("Core: ")
        bold_run.bold = True
        para.add_run(", ".join(skills))

    # --- Experience ---
    doc.add_paragraph("Experience", style="Heading 2")
    for position in resume_data.get("positions", []):
        title = position.get("title", "")
        company = position.get("company", "")
        dates = position.get("dates", "")

        # "Title -- Company" in bold
        para = doc.add_paragraph(style="Normal")
        bold_run = para.add_run(f"{title} -- {company}")
        bold_run.bold = True

        # Dates in italic on separate paragraph
        if dates:
            date_para = doc.add_paragraph(style="Normal")
            italic_run = date_para.add_run(dates)
            italic_run.italic = True

        # Achievement bullets
        for achievement in position.get("achievements", []):
            doc.add_paragraph(achievement, style="List Bullet")

    # --- Education ---
    doc.add_paragraph("Education", style="Heading 2")
    for edu in resume_data.get("education", []):
        degree = edu.get("degree", "")
        institution = edu.get("institution", "")
        year = edu.get("year", "")
        entry = f"{degree} -- {institution}"
        if year:
            entry += f", {year}"
        doc.add_paragraph(entry, style="Normal")

    # Serialize to BytesIO
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
