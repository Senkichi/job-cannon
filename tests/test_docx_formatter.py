"""Tests for job_finder.web.docx_formatter module.

Covers:
- build_resume_docx returns a BytesIO at position 0
- Output is valid .docx (can be opened by python-docx Document)
- Name and contact_line appear in the document
- Professional Summary section is present
- Technical Skills section is present with skill text
- Experience section contains position titles and companies
- Education section contains degree entries
- Missing optional fields do not raise exceptions
- Empty positions/education lists produce no crash
"""

import io

import pytest


def _make_resume_data(**overrides) -> dict:
    """Build a minimal valid resume_data dict for testing."""
    base = {
        "name": "Jane Smith",
        "contact_line": "jane@example.com | 555-1234 | linkedin.com/in/jane",
        "summary": "Experienced software engineer with 10 years in Python.",
        "skills": ["Python", "SQL", "Flask"],
        "positions": [
            {
                "title": "Senior Engineer",
                "company": "Acme Corp",
                "dates": "2020 - Present",
                "achievements": [
                    "Led migration of legacy systems to cloud.",
                    "Reduced API latency by 40%.",
                ],
            }
        ],
        "education": [
            {
                "degree": "B.S. Computer Science",
                "institution": "State University",
                "year": "2012",
            }
        ],
    }
    base.update(overrides)
    return base


class TestBuildResumeDocxOutput:
    """Tests for the output type and structure of build_resume_docx."""

    def test_returns_bytes_io(self):
        """build_resume_docx returns a BytesIO object."""
        from job_finder.web.docx_formatter import build_resume_docx

        result = build_resume_docx(_make_resume_data())
        assert isinstance(result, io.BytesIO)

    def test_buffer_at_position_zero(self):
        """Returned BytesIO is seeked to position 0."""
        from job_finder.web.docx_formatter import build_resume_docx

        result = build_resume_docx(_make_resume_data())
        assert result.tell() == 0

    def test_buffer_is_nonempty(self):
        """Returned BytesIO contains data."""
        from job_finder.web.docx_formatter import build_resume_docx

        result = build_resume_docx(_make_resume_data())
        content = result.read()
        assert len(content) > 0

    def test_output_is_valid_docx(self):
        """Output can be opened as a valid Word document by python-docx."""
        from docx import Document
        from job_finder.web.docx_formatter import build_resume_docx

        buffer = build_resume_docx(_make_resume_data())
        # Should not raise
        doc = Document(buffer)
        assert doc is not None


class TestBuildResumeDocxContent:
    """Tests for the textual content of the generated .docx."""

    def _get_full_text(self, buffer) -> str:
        """Extract all paragraph text from a docx buffer."""
        from docx import Document

        doc = Document(buffer)
        return "\n".join(p.text for p in doc.paragraphs)

    def test_name_appears_in_document(self):
        """Candidate name appears in the document."""
        from job_finder.web.docx_formatter import build_resume_docx

        buffer = build_resume_docx(_make_resume_data(name="Alice Johnson"))
        text = self._get_full_text(buffer)
        assert "Alice Johnson" in text

    def test_contact_line_appears_in_document(self):
        """Contact line appears in the document."""
        from job_finder.web.docx_formatter import build_resume_docx

        buffer = build_resume_docx(_make_resume_data())
        text = self._get_full_text(buffer)
        assert "jane@example.com" in text

    def test_summary_section_present(self):
        """Professional Summary heading and text appear in document."""
        from job_finder.web.docx_formatter import build_resume_docx

        buffer = build_resume_docx(_make_resume_data())
        text = self._get_full_text(buffer)
        assert "Professional Summary" in text
        assert "Experienced software engineer" in text

    def test_skills_section_present(self):
        """Technical Skills section appears with skill keywords."""
        from job_finder.web.docx_formatter import build_resume_docx

        buffer = build_resume_docx(_make_resume_data())
        text = self._get_full_text(buffer)
        assert "Technical Skills" in text
        assert "Python" in text

    def test_experience_section_has_position_title(self):
        """Experience section includes position title and company."""
        from job_finder.web.docx_formatter import build_resume_docx

        buffer = build_resume_docx(_make_resume_data())
        text = self._get_full_text(buffer)
        assert "Senior Engineer" in text
        assert "Acme Corp" in text

    def test_experience_section_has_achievements(self):
        """Experience section includes achievement bullet text."""
        from job_finder.web.docx_formatter import build_resume_docx

        buffer = build_resume_docx(_make_resume_data())
        text = self._get_full_text(buffer)
        assert "API latency" in text

    def test_education_section_present(self):
        """Education section includes degree and institution."""
        from job_finder.web.docx_formatter import build_resume_docx

        buffer = build_resume_docx(_make_resume_data())
        text = self._get_full_text(buffer)
        assert "Education" in text
        assert "B.S. Computer Science" in text
        assert "State University" in text

    def test_education_year_appears(self):
        """Education year appears in document."""
        from job_finder.web.docx_formatter import build_resume_docx

        buffer = build_resume_docx(_make_resume_data())
        text = self._get_full_text(buffer)
        assert "2012" in text


class TestBuildResumeDocxEdgeCases:
    """Tests for missing/empty fields."""

    def test_empty_name_does_not_raise(self):
        """Empty name string does not raise an exception."""
        from job_finder.web.docx_formatter import build_resume_docx

        buffer = build_resume_docx(_make_resume_data(name=""))
        assert buffer.tell() == 0

    def test_empty_positions_does_not_raise(self):
        """Empty positions list does not raise an exception."""
        from job_finder.web.docx_formatter import build_resume_docx

        buffer = build_resume_docx(_make_resume_data(positions=[]))
        assert buffer.tell() == 0

    def test_empty_education_does_not_raise(self):
        """Empty education list does not raise an exception."""
        from job_finder.web.docx_formatter import build_resume_docx

        buffer = build_resume_docx(_make_resume_data(education=[]))
        assert buffer.tell() == 0

    def test_empty_skills_does_not_raise(self):
        """Empty skills list does not raise an exception."""
        from job_finder.web.docx_formatter import build_resume_docx

        buffer = build_resume_docx(_make_resume_data(skills=[]))
        assert buffer.tell() == 0

    def test_missing_contact_line_omitted(self):
        """Missing contact_line key is handled without crash."""
        from job_finder.web.docx_formatter import build_resume_docx

        data = _make_resume_data()
        del data["contact_line"]
        buffer = build_resume_docx(data)
        assert buffer.tell() == 0

    def test_position_without_dates_does_not_crash(self):
        """Position dict without 'dates' key is handled safely."""
        from job_finder.web.docx_formatter import build_resume_docx

        data = _make_resume_data(positions=[{
            "title": "Engineer",
            "company": "Corp",
            "achievements": ["Did things."],
        }])
        buffer = build_resume_docx(data)
        assert buffer.tell() == 0

    def test_education_without_year_does_not_crash(self):
        """Education dict without 'year' key is handled safely."""
        from job_finder.web.docx_formatter import build_resume_docx

        data = _make_resume_data(education=[{
            "degree": "B.S.",
            "institution": "University",
        }])
        buffer = build_resume_docx(data)
        assert buffer.tell() == 0
