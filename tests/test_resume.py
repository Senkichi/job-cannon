"""Tests for Phase 4 resume generation infrastructure.

Covers:
- Migration 4: extends resume_generations table with status tracking columns
- docx_formatter: build_resume_docx produces valid .docx BytesIO
- drive_uploader: upload_to_drive handles .docx vs Google Doc modes
- drive_uploader: get_drive_service detects missing drive.file scope
- resume_generator: generate_resume_single and _generate_resume_background
- resume blueprint: POST /generate and GET /status routes
"""

import io
import sqlite3
from unittest.mock import MagicMock, patch

import pytest


class TestDocxFormatter:
    """build_resume_docx produces a valid .docx BytesIO."""

    def test_returns_bytesio(self, sample_resume_data):
        """build_resume_docx returns an io.BytesIO object."""
        from job_finder.web.docx_formatter import build_resume_docx

        result = build_resume_docx(sample_resume_data)
        assert isinstance(result, io.BytesIO), f"Expected BytesIO, got: {type(result)}"

    def test_bytesio_is_non_empty(self, sample_resume_data):
        """Returned BytesIO has non-zero length (is a real document)."""
        from job_finder.web.docx_formatter import build_resume_docx

        result = build_resume_docx(sample_resume_data)
        data = result.read()
        assert len(data) > 0, "build_resume_docx returned empty BytesIO"

    def test_produces_valid_docx(self, sample_resume_data):
        """BytesIO can be re-opened as a python-docx Document (valid .docx)."""
        from docx import Document

        from job_finder.web.docx_formatter import build_resume_docx

        result = build_resume_docx(sample_resume_data)
        # Should not raise
        doc = Document(result)
        assert doc is not None

    def test_contains_candidate_name(self, sample_resume_data):
        """Document contains the candidate's name."""
        from docx import Document

        from job_finder.web.docx_formatter import build_resume_docx

        result = build_resume_docx(sample_resume_data)
        doc = Document(result)
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert sample_resume_data["name"] in full_text, (
            f"Name '{sample_resume_data['name']}' not found in document"
        )

    def test_contains_summary(self, sample_resume_data):
        """Document contains the professional summary text."""
        from docx import Document

        from job_finder.web.docx_formatter import build_resume_docx

        result = build_resume_docx(sample_resume_data)
        doc = Document(result)
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert sample_resume_data["summary"] in full_text, (
            "Summary text not found in document"
        )

    def test_contains_skills(self, sample_resume_data):
        """Document contains at least one skill from the skills list."""
        from docx import Document

        from job_finder.web.docx_formatter import build_resume_docx

        result = build_resume_docx(sample_resume_data)
        doc = Document(result)
        full_text = "\n".join(p.text for p in doc.paragraphs)
        for skill in sample_resume_data["skills"]:
            assert skill in full_text, f"Skill '{skill}' not found in document"

    def test_contains_position_title(self, sample_resume_data):
        """Document contains job titles from experience."""
        from docx import Document

        from job_finder.web.docx_formatter import build_resume_docx

        result = build_resume_docx(sample_resume_data)
        doc = Document(result)
        full_text = "\n".join(p.text for p in doc.paragraphs)
        first_position = sample_resume_data["positions"][0]
        assert first_position["title"] in full_text, (
            f"Position title '{first_position['title']}' not found in document"
        )

    def test_contains_education(self, sample_resume_data):
        """Document contains education institution."""
        from docx import Document

        from job_finder.web.docx_formatter import build_resume_docx

        result = build_resume_docx(sample_resume_data)
        doc = Document(result)
        full_text = "\n".join(p.text for p in doc.paragraphs)
        edu = sample_resume_data["education"][0]
        assert edu["institution"] in full_text, (
            f"Institution '{edu['institution']}' not found in document"
        )


class TestDriveUpload:
    """upload_to_drive handles .docx and Google Doc conversion modes."""

    def _make_mock_service(self, file_id="abc123", web_view_link="https://docs.google.com/doc/abc123"):
        """Build a mock Drive service that returns a file create response."""
        mock_service = MagicMock()
        mock_service.files.return_value.create.return_value.execute.return_value = {
            "id": file_id,
            "webViewLink": web_view_link,
        }
        return mock_service

    def test_returns_web_view_link(self):
        """upload_to_drive returns the webViewLink from the API response."""
        from job_finder.web.drive_uploader import upload_to_drive

        service = self._make_mock_service(web_view_link="https://docs.google.com/doc/abc123")
        buffer = io.BytesIO(b"fake docx content")

        result = upload_to_drive(service, "Test Resume", buffer, folder_id="folder1")

        assert result == "https://docs.google.com/doc/abc123", (
            f"Expected webViewLink, got: {result}"
        )

    def test_convert_to_gdoc_sets_mime_type(self):
        """When convert_to_gdoc=True, file_metadata includes Google Docs mimeType."""
        from job_finder.web.drive_uploader import upload_to_drive

        service = self._make_mock_service()
        buffer = io.BytesIO(b"fake docx content")

        upload_to_drive(service, "Test Resume", buffer, folder_id="folder1", convert_to_gdoc=True)

        # Inspect the call arguments
        call_kwargs = service.files.return_value.create.call_args
        body_arg = call_kwargs[1].get("body") or call_kwargs[0][0] if call_kwargs[0] else None
        if body_arg is None:
            # Try keyword arg
            body_arg = call_kwargs.kwargs.get("body")

        assert body_arg is not None, "No body argument passed to files().create()"
        assert body_arg.get("mimeType") == "application/vnd.google-apps.document", (
            f"Expected Google Docs mimeType, got: {body_arg.get('mimeType')}"
        )

    def test_no_conversion_omits_mime_type(self):
        """When convert_to_gdoc=False, file_metadata does NOT include mimeType."""
        from job_finder.web.drive_uploader import upload_to_drive

        service = self._make_mock_service()
        buffer = io.BytesIO(b"fake docx content")

        upload_to_drive(service, "Test Resume", buffer, folder_id="folder1", convert_to_gdoc=False)

        call_kwargs = service.files.return_value.create.call_args
        body_arg = call_kwargs[1].get("body") or call_kwargs[0][0] if call_kwargs[0] else None
        if body_arg is None:
            body_arg = call_kwargs.kwargs.get("body")

        assert body_arg is not None, "No body argument passed to files().create()"
        assert "mimeType" not in body_arg, (
            f"Expected no mimeType in metadata, but found: {body_arg.get('mimeType')}"
        )

    def test_no_conversion_appends_docx_extension(self):
        """When convert_to_gdoc=False, .docx is appended to the document name."""
        from job_finder.web.drive_uploader import upload_to_drive

        service = self._make_mock_service()
        buffer = io.BytesIO(b"fake docx content")

        upload_to_drive(service, "Test Resume", buffer, folder_id="folder1", convert_to_gdoc=False)

        call_kwargs = service.files.return_value.create.call_args
        body_arg = call_kwargs[1].get("body") or call_kwargs[0][0] if call_kwargs[0] else None
        if body_arg is None:
            body_arg = call_kwargs.kwargs.get("body")

        assert body_arg is not None, "No body argument passed to files().create()"
        assert body_arg["name"].endswith(".docx"), (
            f"Expected name to end with .docx, got: {body_arg['name']}"
        )

    def test_fallback_url_when_no_web_view_link(self):
        """upload_to_drive constructs fallback URL if webViewLink absent from response."""
        from job_finder.web.drive_uploader import upload_to_drive

        service = MagicMock()
        service.files.return_value.create.return_value.execute.return_value = {
            "id": "fileid123",
            # no webViewLink
        }
        buffer = io.BytesIO(b"fake docx content")

        result = upload_to_drive(service, "Test Resume", buffer, folder_id="folder1")

        assert "fileid123" in result, (
            f"Expected fallback URL containing file ID, got: {result}"
        )


class TestDriveServiceScopeCheck:
    """get_drive_service detects missing drive.file scope and raises ValueError."""

    def test_raises_value_error_when_drive_scope_missing(self, tmp_path):
        """get_drive_service raises ValueError if token lacks drive.file scope.
        Mock reflects scopes=None loading: creds.scopes shows actual granted scopes.
        """
        from job_finder.web.drive_uploader import get_drive_service

        mock_creds = MagicMock()
        # scopes=None loading: creds.scopes reflects what was actually granted
        mock_creds.scopes = ["https://www.googleapis.com/auth/gmail.readonly"]
        mock_creds.valid = True
        mock_creds.expired = False

        token_path = str(tmp_path / "token.json")
        # Write a fake token file so the path exists
        with open(token_path, "w") as f:
            f.write('{"token": "fake"}')

        with patch(
            "job_finder.web.drive_uploader.Credentials.from_authorized_user_file",
            return_value=mock_creds,
        ) as mock_load:
            with pytest.raises(ValueError, match="drive.file"):
                get_drive_service(token_path=token_path)
            # Verify called with scopes=None for honest scope detection
            mock_load.assert_called_once_with(token_path, scopes=None)

    def test_raises_value_error_message_contains_gmail_auth_command(self, tmp_path):
        """ValueError message instructs user to run python -m job_finder.gmail_auth."""
        from job_finder.web.drive_uploader import get_drive_service

        mock_creds = MagicMock()
        mock_creds.scopes = ["https://www.googleapis.com/auth/gmail.readonly"]
        mock_creds.valid = True
        mock_creds.expired = False

        token_path = str(tmp_path / "token.json")
        with open(token_path, "w") as f:
            f.write('{"token": "fake"}')

        with patch(
            "job_finder.web.drive_uploader.Credentials.from_authorized_user_file",
            return_value=mock_creds,
        ):
            with pytest.raises(ValueError) as exc_info:
                get_drive_service(token_path=token_path)

        assert "python -m job_finder.gmail_auth" in str(exc_info.value), (
            f"Error message should contain 'python -m job_finder.gmail_auth', got: {exc_info.value}"
        )

    def test_raises_file_not_found_when_token_missing(self, tmp_path):
        """get_drive_service raises FileNotFoundError when token_path does not exist."""
        from job_finder.web.drive_uploader import get_drive_service

        token_path = str(tmp_path / "no_token.json")  # does not exist

        with pytest.raises(FileNotFoundError):
            get_drive_service(token_path=token_path)

    def test_succeeds_when_drive_scope_present(self, tmp_path):
        """get_drive_service returns service when token includes drive.file scope."""
        from job_finder.web.drive_uploader import get_drive_service

        mock_creds = MagicMock()
        mock_creds.scopes = [
            "https://www.googleapis.com/auth/gmail.readonly",
            "https://www.googleapis.com/auth/drive.file",
        ]
        mock_creds.valid = True
        mock_creds.expired = False

        token_path = str(tmp_path / "token.json")
        with open(token_path, "w") as f:
            f.write('{"token": "fake"}')

        mock_service = MagicMock()
        with patch(
            "job_finder.web.drive_uploader.Credentials.from_authorized_user_file",
            return_value=mock_creds,
        ), patch(
            "job_finder.web.drive_uploader.build",
            return_value=mock_service,
        ):
            result = get_drive_service(token_path=token_path)

        assert result is mock_service

    def test_refreshes_expired_token(self, tmp_path):
        """get_drive_service calls creds.refresh() when token is expired."""
        from job_finder.web.drive_uploader import get_drive_service

        mock_creds = MagicMock()
        mock_creds.scopes = [
            "https://www.googleapis.com/auth/gmail.readonly",
            "https://www.googleapis.com/auth/drive.file",
        ]
        mock_creds.valid = False
        mock_creds.expired = True
        mock_creds.refresh_token = "fake_refresh"

        token_path = str(tmp_path / "token.json")
        with open(token_path, "w") as f:
            f.write('{"token": "fake"}')

        mock_service = MagicMock()
        with patch(
            "job_finder.web.drive_uploader.Credentials.from_authorized_user_file",
            return_value=mock_creds,
        ), patch(
            "job_finder.web.drive_uploader.build",
            return_value=mock_service,
        ), patch(
            "job_finder.web.drive_uploader.Request",
        ) as mock_request:
            result = get_drive_service(token_path=token_path)

        mock_creds.refresh.assert_called_once()
        assert result is mock_service

    def test_raises_value_error_with_actionable_message_when_refresh_fails(self, tmp_path):
        """get_drive_service raises ValueError with actionable message when refresh fails."""
        from google.auth.exceptions import TransportError
        from job_finder.web.drive_uploader import get_drive_service

        mock_creds = MagicMock()
        mock_creds.scopes = [
            "https://www.googleapis.com/auth/gmail.readonly",
            "https://www.googleapis.com/auth/drive.file",
        ]
        mock_creds.valid = False
        mock_creds.expired = True
        mock_creds.refresh_token = "fake_refresh"
        mock_creds.refresh.side_effect = TransportError("network error")

        token_path = str(tmp_path / "token.json")
        with open(token_path, "w") as f:
            f.write('{"token": "fake"}')

        with patch(
            "job_finder.web.drive_uploader.Credentials.from_authorized_user_file",
            return_value=mock_creds,
        ), patch(
            "job_finder.web.drive_uploader.Request",
        ):
            with pytest.raises(ValueError) as exc_info:
                get_drive_service(token_path=token_path)

        assert "python -m job_finder.gmail_auth" in str(exc_info.value), (
            f"Error should contain re-auth command, got: {exc_info.value}"
        )


# =============================================================================
# Task 1 Tests: drive_status.py -- get_drive_status helper
# =============================================================================


class TestDriveStatus:
    """get_drive_status returns structured dict for all failure modes + happy path."""

    def _make_app(self):
        """Create minimal Flask app for request context."""
        from flask import Flask
        app = Flask(__name__)
        app.config["TESTING"] = True
        return app

    def test_returns_no_token_when_token_missing(self, tmp_path):
        """get_drive_status returns error_code='no_token' when token.json absent."""
        from job_finder.web.drive_status import get_drive_status

        config = {"drive": {"folder_id": "some-folder"}}
        token_path = str(tmp_path / "no_token.json")

        app = self._make_app()
        with app.test_request_context():
            result = get_drive_status(config, token_path=token_path)

        assert result["ok"] is False
        assert result["error_code"] == "no_token"

    def test_returns_missing_scope_when_drive_file_absent(self, tmp_path):
        """get_drive_status returns error_code='missing_scope' when token lacks drive.file."""
        from job_finder.web.drive_status import get_drive_status

        token_path = str(tmp_path / "token.json")
        with open(token_path, "w") as f:
            f.write('{"token": "fake"}')

        config = {"drive": {"folder_id": "some-folder"}}

        mock_creds = MagicMock()
        mock_creds.scopes = ["https://www.googleapis.com/auth/gmail.readonly"]
        mock_creds.expired = False

        app = self._make_app()
        with app.test_request_context():
            with patch(
                "google.oauth2.credentials.Credentials.from_authorized_user_file",
                return_value=mock_creds,
            ):
                result = get_drive_status(config, token_path=token_path)

        assert result["ok"] is False
        assert result["error_code"] == "missing_scope"

    def test_returns_refresh_failed_when_token_refresh_raises(self, tmp_path):
        """get_drive_status returns error_code='refresh_failed' when refresh raises."""
        from google.auth.exceptions import TransportError
        from job_finder.web.drive_status import get_drive_status

        token_path = str(tmp_path / "token.json")
        with open(token_path, "w") as f:
            f.write('{"token": "fake"}')

        config = {"drive": {"folder_id": "some-folder"}}

        mock_creds = MagicMock()
        mock_creds.scopes = [
            "https://www.googleapis.com/auth/gmail.readonly",
            "https://www.googleapis.com/auth/drive.file",
        ]
        mock_creds.expired = True
        mock_creds.refresh_token = "fake_refresh"
        mock_creds.refresh.side_effect = TransportError("network error")

        app = self._make_app()
        with app.test_request_context():
            with patch(
                "google.oauth2.credentials.Credentials.from_authorized_user_file",
                return_value=mock_creds,
            ), patch("google.auth.transport.requests.Request"):
                result = get_drive_status(config, token_path=token_path)

        assert result["ok"] is False
        assert result["error_code"] == "refresh_failed"

    def test_returns_no_folder_id_when_folder_id_empty(self, tmp_path):
        """get_drive_status returns error_code='no_folder_id' when folder_id is empty."""
        from job_finder.web.drive_status import get_drive_status

        token_path = str(tmp_path / "token.json")
        with open(token_path, "w") as f:
            f.write('{"token": "fake"}')

        config = {"drive": {"folder_id": ""}}  # empty folder_id

        mock_creds = MagicMock()
        mock_creds.scopes = [
            "https://www.googleapis.com/auth/gmail.readonly",
            "https://www.googleapis.com/auth/drive.file",
        ]
        mock_creds.expired = False

        app = self._make_app()
        with app.test_request_context():
            with patch(
                "google.oauth2.credentials.Credentials.from_authorized_user_file",
                return_value=mock_creds,
            ):
                result = get_drive_status(config, token_path=token_path)

        assert result["ok"] is False
        assert result["error_code"] == "no_folder_id"

    def test_returns_ok_true_when_all_conditions_met(self, tmp_path):
        """get_drive_status returns ok=True when token has drive.file scope and folder_id set."""
        from job_finder.web.drive_status import get_drive_status

        token_path = str(tmp_path / "token.json")
        with open(token_path, "w") as f:
            f.write('{"token": "fake"}')

        config = {"drive": {"folder_id": "real-folder-id"}}

        mock_creds = MagicMock()
        mock_creds.scopes = [
            "https://www.googleapis.com/auth/gmail.readonly",
            "https://www.googleapis.com/auth/drive.file",
        ]
        mock_creds.expired = False

        app = self._make_app()
        with app.test_request_context():
            with patch(
                "google.oauth2.credentials.Credentials.from_authorized_user_file",
                return_value=mock_creds,
            ):
                result = get_drive_status(config, token_path=token_path)

        assert result["ok"] is True
        assert result.get("error_code") is None

    def test_caches_result_on_flask_g(self, tmp_path):
        """get_drive_status caches on Flask g; second call returns same object."""
        from job_finder.web.drive_status import get_drive_status

        token_path = str(tmp_path / "token.json")
        with open(token_path, "w") as f:
            f.write('{"token": "fake"}')

        config = {"drive": {"folder_id": "real-folder-id"}}

        mock_creds = MagicMock()
        mock_creds.scopes = [
            "https://www.googleapis.com/auth/gmail.readonly",
            "https://www.googleapis.com/auth/drive.file",
        ]
        mock_creds.expired = False

        app = self._make_app()
        with app.test_request_context():
            with patch(
                "google.oauth2.credentials.Credentials.from_authorized_user_file",
                return_value=mock_creds,
            ) as mock_load:
                result1 = get_drive_status(config, token_path=token_path)
                result2 = get_drive_status(config, token_path=token_path)

            # Called only once — second call uses cached g.drive_status
            assert mock_load.call_count == 1
            assert result1 is result2


# =============================================================================
# Task 1 Tests: resume_generator.py -- generate_resume_single + background
# =============================================================================

class TestSinglePassGeneration:
    """generate_resume_single returns structured resume dict via Sonnet."""

    def test_returns_resume_dict(self, tmp_db_path, sample_resume_data):
        """generate_resume_single returns a dict with expected resume keys."""
        from job_finder.web.db_migrate import run_migrations
        from job_finder.web.resume_generator import generate_resume_single

        run_migrations(tmp_db_path)
        conn = __import__("sqlite3").connect(tmp_db_path)

        mock_client = MagicMock()
        mock_response = MagicMock()
        mock_response.content = [MagicMock()]
        mock_response.content[0].input = {
            "name": "Jane Doe",
            "contact_line": "jane@example.com",
            "summary": "Experienced data scientist.",
            "skills": ["Python", "SQL"],
            "positions": [
                {
                    "title": "Senior Data Scientist",
                    "company": "Acme Corp",
                    "dates": "Jan 2021 - Present",
                    "achievements": ["Led A/B testing platform"],
                }
            ],
            "education": [{"degree": "M.S. Statistics", "institution": "Stanford", "year": "2018"}],
        }
        mock_response.usage.input_tokens = 100
        mock_response.usage.output_tokens = 200
        mock_client.messages.create.return_value = mock_response

        job_row = {
            "dedup_key": "acme|senior-ds|remote",
            "title": "Senior Data Scientist",
            "company": "Acme Corp",
            "jd_full": "Looking for a senior data scientist with Python and SQL skills.",
            "fit_analysis": None,
        }
        config = {
            "scoring": {"models": {"sonnet": "claude-sonnet-4-6"}, "monthly_budget_usd": 25.0},
        }

        result = generate_resume_single(mock_client, job_row, sample_resume_data, conn, config)
        conn.close()

        assert result is not None
        assert "name" in result
        assert "summary" in result
        assert "skills" in result
        assert "positions" in result

    def test_returns_none_when_budget_exceeded(self, tmp_db_path, sample_resume_data):
        """generate_resume_single returns None when cost_gate returns False."""
        from job_finder.web.db_migrate import run_migrations
        from job_finder.web.resume_generator import generate_resume_single

        run_migrations(tmp_db_path)
        conn = __import__("sqlite3").connect(tmp_db_path)

        # Exhaust the budget by inserting costs exceeding cap
        conn.execute(
            "INSERT INTO scoring_costs (job_id, purpose, model, input_tokens, output_tokens, cost_usd, timestamp) "
            "VALUES (?, ?, ?, ?, ?, ?, ?)",
            (None, "test", "claude-sonnet-4-6", 0, 0, 30.0, "2026-03-01T00:00:00Z"),
        )
        conn.commit()

        mock_client = MagicMock()
        job_row = {
            "dedup_key": "acme|senior-ds|remote",
            "title": "Senior Data Scientist",
            "company": "Acme Corp",
            "jd_full": "Looking for a senior data scientist.",
            "fit_analysis": None,
        }
        config = {
            "scoring": {"models": {"sonnet": "claude-sonnet-4-6"}, "monthly_budget_usd": 25.0},
        }

        result = generate_resume_single(mock_client, job_row, sample_resume_data, conn, config)
        conn.close()

        assert result is None, "Expected None when budget exceeded"

    def test_calls_call_claude_with_resume_generation_purpose(self, tmp_db_path, sample_resume_data):
        """generate_resume_single calls call_claude with purpose='resume_generation'."""
        from job_finder.web.db_migrate import run_migrations
        from job_finder.web.resume_generator import generate_resume_single

        run_migrations(tmp_db_path)
        conn = __import__("sqlite3").connect(tmp_db_path)

        mock_client = MagicMock()
        mock_response = MagicMock()
        mock_response.content = [MagicMock()]
        mock_response.content[0].input = {
            "name": "Jane Doe",
            "contact_line": "jane@example.com",
            "summary": "Experienced data scientist.",
            "skills": ["Python"],
            "positions": [],
            "education": [],
        }
        mock_response.usage.input_tokens = 100
        mock_response.usage.output_tokens = 200
        mock_client.messages.create.return_value = mock_response

        job_row = {
            "dedup_key": "acme|senior-ds|remote",
            "title": "Senior Data Scientist",
            "company": "Acme Corp",
            "jd_full": "Looking for a data scientist.",
            "fit_analysis": None,
        }
        config = {
            "scoring": {"models": {"sonnet": "claude-sonnet-4-6"}, "monthly_budget_usd": 25.0},
        }

        with patch("job_finder.web.resume_generator.call_claude") as mock_call:
            mock_call.return_value = (mock_response.content[0].input, 0.05)
            generate_resume_single(mock_client, job_row, sample_resume_data, conn, config)

        conn.close()

        assert mock_call.called, "call_claude was not called"
        call_kwargs = mock_call.call_args
        # purpose should be 'resume_generation'
        actual_purpose = call_kwargs[1].get("purpose") or call_kwargs[0][7] if call_kwargs[0] else None
        if actual_purpose is None:
            # Try keyword access
            actual_purpose = call_kwargs.kwargs.get("purpose")
        assert actual_purpose == "resume_generation", (
            f"Expected purpose='resume_generation', got: {actual_purpose}"
        )


class TestClosedWorldConstraint:
    """System prompt includes closed-world constraint."""

    def test_system_prompt_contains_closed_world_constraint(self, tmp_db_path, sample_resume_data):
        """generate_resume_single passes system prompt with ONLY/profile constraint text."""
        from job_finder.web.db_migrate import run_migrations
        from job_finder.web.resume_generator import generate_resume_single

        run_migrations(tmp_db_path)
        conn = __import__("sqlite3").connect(tmp_db_path)

        mock_client = MagicMock()
        job_row = {
            "dedup_key": "acme|senior-ds|remote",
            "title": "Senior Data Scientist",
            "company": "Acme Corp",
            "jd_full": "Looking for a data scientist.",
            "fit_analysis": None,
        }
        config = {
            "scoring": {"models": {"sonnet": "claude-sonnet-4-6"}, "monthly_budget_usd": 25.0},
        }

        captured_system = {}
        with patch("job_finder.web.resume_generator.call_claude") as mock_call:
            mock_call.return_value = ({
                "name": "Jane",
                "contact_line": "",
                "summary": "test",
                "skills": [],
                "positions": [],
                "education": [],
            }, 0.01)
            generate_resume_single(mock_client, job_row, sample_resume_data, conn, config)
            if mock_call.called:
                captured_system["system"] = mock_call.call_args.kwargs.get("system") or (
                    mock_call.call_args[0][2] if mock_call.call_args[0] else ""
                )

        conn.close()

        system_prompt = captured_system.get("system", "")
        assert system_prompt, "No system prompt was passed to call_claude"
        # Check for closed-world language -- either ONLY or NEVER invent/add
        lower_system = system_prompt.lower()
        assert "only" in lower_system or "never" in lower_system, (
            "System prompt must contain closed-world constraint language ('ONLY' or 'NEVER')"
        )

    def test_resume_schema_has_required_keys(self):
        """RESUME_SCHEMA has required top-level keys: name, summary, skills, positions."""
        from job_finder.web.resume_generator import RESUME_SCHEMA

        required = RESUME_SCHEMA.get("required", [])
        for key in ["name", "summary", "skills", "positions"]:
            assert key in required, f"RESUME_SCHEMA missing required key: {key}"


class TestGenerationHistoryTracking:
    """_generate_resume_background updates resume_generations status transitions."""

    def test_background_updates_status_to_done(self, tmp_db_path, sample_resume_data):
        """_generate_resume_background sets status='done' and doc_url on success."""
        import sqlite3 as _sqlite3
        from job_finder.web.db_migrate import run_migrations
        from job_finder.web.resume_generator import _generate_resume_background

        run_migrations(tmp_db_path)
        conn = _sqlite3.connect(tmp_db_path)

        # Insert a pending row
        conn.execute(
            "INSERT INTO resume_generations (job_id, generated_at, model, status, generation_type) "
            "VALUES (?, ?, ?, ?, ?)",
            ("acme|senior-ds|remote", "2026-03-11T00:00:00", "claude-sonnet-4-6", "pending", "single"),
        )
        conn.commit()
        gen_id = conn.execute("SELECT last_insert_rowid()").fetchone()[0]
        conn.close()

        job_row = {
            "dedup_key": "acme|senior-ds|remote",
            "title": "Senior Data Scientist",
            "company": "Acme Corp",
            "jd_full": "Looking for a data scientist.",
            "fit_analysis": None,
        }
        config = {
            "scoring": {"models": {"sonnet": "claude-sonnet-4-6"}, "monthly_budget_usd": 25.0},
            "drive": {"folder_id": "test-folder", "convert_to_gdoc": True},
        }

        mock_resume_data = {
            "name": "Jane Doe",
            "contact_line": "jane@example.com",
            "summary": "Experienced scientist.",
            "skills": ["Python"],
            "positions": [],
            "education": [],
        }

        with patch("job_finder.web.resume_generator.generate_resume_single", return_value=mock_resume_data):
            with patch("job_finder.web.resume_generator.build_resume_docx") as mock_docx:
                mock_docx.return_value = __import__("io").BytesIO(b"fake-docx")
                with patch("job_finder.web.resume_generator.get_drive_service") as mock_svc:
                    with patch("job_finder.web.resume_generator.upload_to_drive", return_value="https://docs.google.com/doc/xyz"):
                        _generate_resume_background(tmp_db_path, gen_id, job_row, sample_resume_data, config)

        verify_conn = _sqlite3.connect(tmp_db_path)
        row = verify_conn.execute(
            "SELECT status, doc_url FROM resume_generations WHERE id = ?", (gen_id,)
        ).fetchone()
        verify_conn.close()

        assert row[0] == "done", f"Expected status='done', got: {row[0]}"
        assert row[1] == "https://docs.google.com/doc/xyz", f"Expected doc_url set, got: {row[1]}"

    def test_background_sets_status_error_on_exception(self, tmp_db_path, sample_resume_data):
        """_generate_resume_background sets status='error' and error_msg on exception."""
        import sqlite3 as _sqlite3
        from job_finder.web.db_migrate import run_migrations
        from job_finder.web.resume_generator import _generate_resume_background

        run_migrations(tmp_db_path)
        conn = _sqlite3.connect(tmp_db_path)

        conn.execute(
            "INSERT INTO resume_generations (job_id, generated_at, model, status, generation_type) "
            "VALUES (?, ?, ?, ?, ?)",
            ("acme|senior-ds|remote", "2026-03-11T00:00:00", "claude-sonnet-4-6", "pending", "single"),
        )
        conn.commit()
        gen_id = conn.execute("SELECT last_insert_rowid()").fetchone()[0]
        conn.close()

        job_row = {
            "dedup_key": "acme|senior-ds|remote",
            "title": "Senior Data Scientist",
            "company": "Acme Corp",
            "jd_full": "Looking for a data scientist.",
            "fit_analysis": None,
        }
        config = {
            "scoring": {"models": {"sonnet": "claude-sonnet-4-6"}, "monthly_budget_usd": 25.0},
            "drive": {"folder_id": "test-folder"},
        }

        with patch("job_finder.web.resume_generator.generate_resume_single", side_effect=RuntimeError("Simulated failure")):
            _generate_resume_background(tmp_db_path, gen_id, job_row, sample_resume_data, config)

        verify_conn = _sqlite3.connect(tmp_db_path)
        row = verify_conn.execute(
            "SELECT status, error_msg FROM resume_generations WHERE id = ?", (gen_id,)
        ).fetchone()
        verify_conn.close()

        assert row[0] == "error", f"Expected status='error', got: {row[0]}"
        assert "Simulated failure" in row[1], f"Expected error_msg to contain error text, got: {row[1]}"

    def test_background_sets_error_when_budget_exceeded(self, tmp_db_path, sample_resume_data):
        """_generate_resume_background sets status='error' with budget msg when generate_resume_single returns None."""
        import sqlite3 as _sqlite3
        from job_finder.web.db_migrate import run_migrations
        from job_finder.web.resume_generator import _generate_resume_background

        run_migrations(tmp_db_path)
        conn = _sqlite3.connect(tmp_db_path)

        conn.execute(
            "INSERT INTO resume_generations (job_id, generated_at, model, status, generation_type) "
            "VALUES (?, ?, ?, ?, ?)",
            ("acme|senior-ds|remote", "2026-03-11T00:00:00", "claude-sonnet-4-6", "pending", "single"),
        )
        conn.commit()
        gen_id = conn.execute("SELECT last_insert_rowid()").fetchone()[0]
        conn.close()

        job_row = {
            "dedup_key": "acme|senior-ds|remote",
            "title": "Senior Data Scientist",
            "company": "Acme Corp",
            "jd_full": "Looking for a data scientist.",
            "fit_analysis": None,
        }
        config = {
            "scoring": {"models": {"sonnet": "claude-sonnet-4-6"}, "monthly_budget_usd": 25.0},
            "drive": {"folder_id": "test-folder"},
        }

        with patch("job_finder.web.resume_generator.generate_resume_single", return_value=None):
            _generate_resume_background(tmp_db_path, gen_id, job_row, sample_resume_data, config)

        verify_conn = _sqlite3.connect(tmp_db_path)
        row = verify_conn.execute(
            "SELECT status, error_msg FROM resume_generations WHERE id = ?", (gen_id,)
        ).fetchone()
        verify_conn.close()

        assert row[0] == "error", f"Expected status='error', got: {row[0]}"
        assert row[1] is not None, "Expected error_msg to be set"


# =============================================================================
# Task 2 Tests: resume blueprint routes (POST /generate, GET /status)
# =============================================================================

class TestResumeRoutes:
    """Resume blueprint: generate and status polling routes."""

    @pytest.fixture
    def app_with_resume_bp(self, tmp_db_path):
        """Create test app with resume blueprint registered and a Sonnet-scored job."""
        import sqlite3 as _sqlite3

        from job_finder.web.db_migrate import run_migrations
        from job_finder.web import create_app

        run_migrations(tmp_db_path)

        # Insert a Sonnet-scored job
        conn = _sqlite3.connect(tmp_db_path)
        conn.execute(
            """INSERT INTO jobs
                (dedup_key, title, company, location, sources, source_urls,
                 source_id, first_seen, last_seen, score, score_breakdown,
                 user_interest, pipeline_status, sonnet_score, jd_full)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
            (
                "acme|senior-ds|remote",
                "Senior Data Scientist",
                "Acme Corp",
                "Remote",
                '["linkedin"]',
                '["https://linkedin.com/jobs/1234"]',
                "1234",
                "2026-03-01T00:00:00",
                "2026-03-11T00:00:00",
                8.5,
                "{}",
                "reviewing",
                "reviewing",
                85.0,
                "Full job description text here.",
            ),
        )
        conn.commit()
        conn.close()

        cfg = {
            "db": {"path": tmp_db_path},
            "scoring": {"models": {"sonnet": "claude-sonnet-4-6"}, "monthly_budget_usd": 25.0},
            "drive": {"folder_id": "test-folder", "convert_to_gdoc": True},
        }
        app = create_app(config=cfg)
        app.config["TESTING"] = True
        return app

    def test_generate_returns_200_with_polling_fragment(self, app_with_resume_bp):
        """POST /jobs/<key>/resume/generate returns 200 with polling fragment."""
        app = app_with_resume_bp
        dedup_key = "acme|senior-ds|remote"
        from urllib.parse import quote

        with patch("job_finder.web.blueprints.resume._generate_resume_background"):
            with patch("job_finder.web.blueprints.resume.load_profile", return_value={}):
                with patch("threading.Thread") as mock_thread:
                    mock_thread.return_value.start = MagicMock()
                    with app.test_client() as client:
                        resp = client.post(
                            f"/jobs/{quote(dedup_key, safe='')}/resume/generate",
                        )

        assert resp.status_code == 200, f"Expected 200, got: {resp.status_code}"
        body = resp.data.decode()
        # Should contain HTMX polling attributes
        assert "hx-get" in body or "hx-trigger" in body or "Generating" in body, (
            "Response should contain polling fragment"
        )

    def test_generate_returns_400_when_no_sonnet_score(self, tmp_db_path):
        """POST /jobs/<key>/resume/generate returns 400 when job has no sonnet_score."""
        import sqlite3 as _sqlite3
        from urllib.parse import quote

        from job_finder.web.db_migrate import run_migrations
        from job_finder.web import create_app

        run_migrations(tmp_db_path)

        # Insert a Haiku-only job (no sonnet_score)
        conn = _sqlite3.connect(tmp_db_path)
        conn.execute(
            """INSERT INTO jobs
                (dedup_key, title, company, location, sources, source_urls,
                 source_id, first_seen, last_seen, score, score_breakdown,
                 user_interest, pipeline_status, haiku_score)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
            (
                "haiku|only|job",
                "Data Analyst",
                "Some Corp",
                "Remote",
                '["linkedin"]',
                '["https://linkedin.com/jobs/5678"]',
                "5678",
                "2026-03-01T00:00:00",
                "2026-03-11T00:00:00",
                6.0,
                "{}",
                "reviewing",
                "reviewing",
                65.0,
            ),
        )
        conn.commit()
        conn.close()

        cfg = {
            "db": {"path": tmp_db_path},
            "scoring": {"models": {"sonnet": "claude-sonnet-4-6"}, "monthly_budget_usd": 25.0},
        }
        app = create_app(config=cfg)
        app.config["TESTING"] = True
        dedup_key = "haiku|only|job"

        with app.test_client() as client:
            resp = client.post(f"/jobs/{quote(dedup_key, safe='')}/resume/generate")

        assert resp.status_code == 400, f"Expected 400, got: {resp.status_code}"

    def test_status_returns_done_template(self, app_with_resume_bp, tmp_db_path):
        """GET /jobs/<key>/resume/status/<id> returns done template when status='done'."""
        import sqlite3 as _sqlite3
        from urllib.parse import quote

        app = app_with_resume_bp
        dedup_key = "acme|senior-ds|remote"

        # Insert a done generation row
        conn = _sqlite3.connect(tmp_db_path)
        conn.execute(
            "INSERT INTO resume_generations (job_id, generated_at, model, status, doc_url) "
            "VALUES (?, ?, ?, ?, ?)",
            (dedup_key, "2026-03-11T00:00:00", "claude-sonnet-4-6", "done", "https://docs.google.com/doc/xyz"),
        )
        conn.commit()
        gen_id = conn.execute("SELECT last_insert_rowid()").fetchone()[0]
        conn.close()

        with app.test_client() as client:
            resp = client.get(f"/jobs/{quote(dedup_key, safe='')}/resume/status/{gen_id}")

        assert resp.status_code == 200
        body = resp.data.decode()
        # Done template has no hx-trigger (stops polling) and has a View Resume link
        assert "View Resume" in body or "docs.google.com" in body, (
            "Done template should contain View Resume link or Drive URL"
        )
        assert "hx-trigger" not in body, "Done template must NOT have hx-trigger (stops polling)"

    def test_status_returns_generating_template_with_polling(self, app_with_resume_bp, tmp_db_path):
        """GET status returns generating template with hx-trigger when status='generating'."""
        import sqlite3 as _sqlite3
        from urllib.parse import quote

        app = app_with_resume_bp
        dedup_key = "acme|senior-ds|remote"

        conn = _sqlite3.connect(tmp_db_path)
        conn.execute(
            "INSERT INTO resume_generations (job_id, generated_at, model, status) "
            "VALUES (?, ?, ?, ?)",
            (dedup_key, "2026-03-11T00:00:00", "claude-sonnet-4-6", "generating"),
        )
        conn.commit()
        gen_id = conn.execute("SELECT last_insert_rowid()").fetchone()[0]
        conn.close()

        with app.test_client() as client:
            resp = client.get(f"/jobs/{quote(dedup_key, safe='')}/resume/status/{gen_id}")

        assert resp.status_code == 200
        body = resp.data.decode()
        assert "hx-trigger" in body, "Generating template must have hx-trigger for polling"

    def test_status_returns_error_template(self, app_with_resume_bp, tmp_db_path):
        """GET status returns error template when status='error'."""
        import sqlite3 as _sqlite3
        from urllib.parse import quote

        app = app_with_resume_bp
        dedup_key = "acme|senior-ds|remote"

        conn = _sqlite3.connect(tmp_db_path)
        conn.execute(
            "INSERT INTO resume_generations (job_id, generated_at, model, status, error_msg) "
            "VALUES (?, ?, ?, ?, ?)",
            (dedup_key, "2026-03-11T00:00:00", "claude-sonnet-4-6", "error", "Budget exceeded"),
        )
        conn.commit()
        gen_id = conn.execute("SELECT last_insert_rowid()").fetchone()[0]
        conn.close()

        with app.test_client() as client:
            resp = client.get(f"/jobs/{quote(dedup_key, safe='')}/resume/status/{gen_id}")

        assert resp.status_code == 200
        body = resp.data.decode()
        assert "Budget exceeded" in body or "error" in body.lower(), (
            "Error template should contain the error message"
        )
        assert "hx-trigger" not in body, "Error template must NOT have hx-trigger (stops polling)"


# =============================================================================
# Plan 03 Tests: Multi-version synthesis
# =============================================================================

class TestMultiVersionStrategySelection:
    """STRATEGY_POOL and _haiku_select_strategies behavior."""

    def test_strategy_pool_has_at_least_5_strategies(self):
        """STRATEGY_POOL contains at least 5 strategy identifiers."""
        from job_finder.web.resume_generator import STRATEGY_POOL

        assert len(STRATEGY_POOL) >= 5, (
            f"Expected at least 5 strategies, got: {len(STRATEGY_POOL)}"
        )

    def test_strategy_pool_contains_strings(self):
        """All items in STRATEGY_POOL are non-empty strings."""
        from job_finder.web.resume_generator import STRATEGY_POOL

        for item in STRATEGY_POOL:
            assert isinstance(item, str) and item, (
                f"STRATEGY_POOL item is not a non-empty string: {item!r}"
            )

    def test_haiku_select_strategies_returns_exactly_3(self, tmp_db_path):
        """_haiku_select_strategies returns exactly 3 strategies."""
        import sqlite3 as _sqlite3
        from unittest.mock import MagicMock, patch

        from job_finder.web.db_migrate import run_migrations
        from job_finder.web.resume_generator import STRATEGY_POOL, _haiku_select_strategies

        run_migrations(tmp_db_path)
        conn = _sqlite3.connect(tmp_db_path)
        mock_client = MagicMock()

        strategies_returned = STRATEGY_POOL[:3]
        config = {
            "scoring": {
                "models": {"haiku": "claude-haiku-4-5"},
                "monthly_budget_usd": 25.0,
            }
        }
        job_row = {
            "dedup_key": "acme|ds|remote",
            "title": "Data Scientist",
            "company": "Acme",
            "jd_full": "Looking for a data scientist.",
        }

        with patch("job_finder.web.resume_generator.call_claude") as mock_call:
            mock_call.return_value = ({"strategies": strategies_returned, "reasoning": "test"}, 0.001)
            result = _haiku_select_strategies(mock_client, job_row, conn, config)

        conn.close()
        assert len(result) == 3, f"Expected 3 strategies, got: {len(result)}"

    def test_haiku_select_strategies_calls_call_claude_with_haiku_model(self, tmp_db_path):
        """_haiku_select_strategies calls call_claude with Haiku model and purpose='resume_strategy'."""
        import sqlite3 as _sqlite3
        from unittest.mock import MagicMock, patch

        from job_finder.web.db_migrate import run_migrations
        from job_finder.web.resume_generator import STRATEGY_POOL, _haiku_select_strategies

        run_migrations(tmp_db_path)
        conn = _sqlite3.connect(tmp_db_path)
        mock_client = MagicMock()

        config = {
            "scoring": {
                "models": {"haiku": "claude-haiku-4-5"},
                "monthly_budget_usd": 25.0,
            }
        }
        job_row = {
            "dedup_key": "acme|ds|remote",
            "title": "Data Scientist",
            "company": "Acme",
            "jd_full": "Looking for a data scientist.",
        }

        with patch("job_finder.web.resume_generator.call_claude") as mock_call:
            mock_call.return_value = ({"strategies": STRATEGY_POOL[:3], "reasoning": "test"}, 0.001)
            _haiku_select_strategies(mock_client, job_row, conn, config)

        conn.close()
        assert mock_call.called, "call_claude was not called"
        call_kwargs = mock_call.call_args.kwargs
        model = call_kwargs.get("model") or mock_call.call_args[0][1]
        purpose = call_kwargs.get("purpose") or mock_call.call_args[0][7]
        assert "haiku" in model.lower(), f"Expected Haiku model, got: {model}"
        assert purpose == "resume_strategy", f"Expected purpose='resume_strategy', got: {purpose}"

    def test_haiku_select_strategies_fallback_on_failure(self, tmp_db_path):
        """_haiku_select_strategies falls back to first 3 from STRATEGY_POOL if call_claude fails."""
        import sqlite3 as _sqlite3
        from unittest.mock import MagicMock, patch

        from job_finder.web.db_migrate import run_migrations
        from job_finder.web.resume_generator import STRATEGY_POOL, _haiku_select_strategies

        run_migrations(tmp_db_path)
        conn = _sqlite3.connect(tmp_db_path)
        mock_client = MagicMock()

        config = {
            "scoring": {
                "models": {"haiku": "claude-haiku-4-5"},
                "monthly_budget_usd": 25.0,
            }
        }
        job_row = {
            "dedup_key": "acme|ds|remote",
            "title": "Data Scientist",
            "company": "Acme",
            "jd_full": "Looking for a data scientist.",
        }

        with patch("job_finder.web.resume_generator.call_claude", side_effect=Exception("API failure")):
            result = _haiku_select_strategies(mock_client, job_row, conn, config)

        conn.close()
        assert result == STRATEGY_POOL[:3], (
            f"Expected fallback to first 3 STRATEGY_POOL items, got: {result}"
        )


class TestParallelVariantGeneration:
    """generate_resume_multi parallel ThreadPoolExecutor behavior."""

    def _make_sample_resume(self):
        return {
            "name": "Jane Doe",
            "contact_line": "jane@example.com",
            "summary": "Experienced scientist.",
            "skills": ["Python", "SQL"],
            "positions": [
                {
                    "title": "Senior DS",
                    "company": "Acme",
                    "dates": "2021-Present",
                    "achievements": ["Did things"],
                }
            ],
            "education": [],
        }

    def test_generate_resume_multi_calls_call_claude_multiple_times(self, tmp_db_path, sample_resume_data):
        """generate_resume_multi invokes call_claude for strategy + 3 variants + synthesis (5+ calls)."""
        import sqlite3 as _sqlite3
        from unittest.mock import MagicMock, patch

        from job_finder.web.db_migrate import run_migrations
        from job_finder.web.resume_generator import STRATEGY_POOL, generate_resume_multi

        run_migrations(tmp_db_path)

        config = {
            "scoring": {
                "models": {
                    "haiku": "claude-haiku-4-5",
                    "sonnet": "claude-sonnet-4-6",
                },
                "monthly_budget_usd": 25.0,
            }
        }
        job_row = {
            "dedup_key": "acme|ds|remote",
            "title": "Data Scientist",
            "company": "Acme",
            "jd_full": "Looking for a data scientist with Python skills.",
            "fit_analysis": None,
            "sonnet_score": 85.0,
        }

        sample_resume = self._make_sample_resume()

        with patch("job_finder.web.resume_generator.call_claude") as mock_call:
            mock_call.return_value = (
                {**sample_resume, "strategies": STRATEGY_POOL[:3], "reasoning": "test"},
                0.01,
            )
            with patch("job_finder.web.resume_generator.cost_gate", return_value=True):
                with patch("job_finder.web.resume_generator.anthropic") as mock_anthropic:
                    mock_anthropic.Anthropic.return_value = MagicMock()
                    generate_resume_multi(tmp_db_path, job_row, sample_resume_data, config)

        # Should be called: 1 (strategy) + 3 (variants) + 1 (synthesis) = 5 times
        assert mock_call.call_count >= 5, (
            f"Expected at least 5 call_claude invocations, got: {mock_call.call_count}"
        )

    def test_generate_resume_multi_uses_different_strategies(self, tmp_db_path, sample_resume_data):
        """Each parallel variant receives a different strategy directive in the system prompt."""
        import sqlite3 as _sqlite3
        from unittest.mock import MagicMock, call, patch

        from job_finder.web.db_migrate import run_migrations
        from job_finder.web.resume_generator import STRATEGY_POOL, generate_resume_multi

        run_migrations(tmp_db_path)

        config = {
            "scoring": {
                "models": {
                    "haiku": "claude-haiku-4-5",
                    "sonnet": "claude-sonnet-4-6",
                },
                "monthly_budget_usd": 25.0,
            }
        }
        job_row = {
            "dedup_key": "acme|ds|remote",
            "title": "Data Scientist",
            "company": "Acme",
            "jd_full": "Looking for a data scientist.",
            "fit_analysis": None,
            "sonnet_score": 85.0,
        }

        sample_resume = self._make_sample_resume()
        system_prompts = []

        original_call_claude = None

        def capturing_call_claude(**kwargs):
            system_prompts.append(kwargs.get("system", ""))
            purpose = kwargs.get("purpose", "")
            if purpose == "resume_strategy":
                return ({"strategies": STRATEGY_POOL[:3], "reasoning": "ok"}, 0.001)
            return (sample_resume, 0.01)

        with patch("job_finder.web.resume_generator.call_claude", side_effect=capturing_call_claude):
            with patch("job_finder.web.resume_generator.cost_gate", return_value=True):
                with patch("job_finder.web.resume_generator.anthropic") as mock_anthropic:
                    mock_anthropic.Anthropic.return_value = MagicMock()
                    generate_resume_multi(tmp_db_path, job_row, sample_resume_data, config)

        # Filter out strategy-selection and synthesis system prompts; keep variant generation prompts
        variant_prompts = [p for p in system_prompts if "STRATEGY EMPHASIS" in p]
        assert len(variant_prompts) == 3, (
            f"Expected 3 variant system prompts with STRATEGY EMPHASIS, got: {len(variant_prompts)}"
        )
        # Each variant prompt should be unique (different strategy)
        assert len(set(variant_prompts)) == 3, (
            "Each variant should use a different strategy (duplicate system prompts found)"
        )


class TestThreadSafety:
    """_generate_single_variant opens its own SQLite connection."""

    def test_each_variant_thread_opens_own_sqlite_connection(self, tmp_db_path, sample_resume_data):
        """_generate_single_variant opens its own sqlite3.connect() (not shared conn)."""
        import sqlite3 as _sqlite3
        from unittest.mock import MagicMock, patch

        from job_finder.web.db_migrate import run_migrations
        from job_finder.web.resume_generator import _generate_single_variant

        run_migrations(tmp_db_path)

        sample_resume = {
            "name": "Jane Doe",
            "contact_line": "",
            "summary": "test",
            "skills": [],
            "positions": [],
            "education": [],
        }

        config = {
            "scoring": {
                "models": {"sonnet": "claude-sonnet-4-6"},
                "monthly_budget_usd": 25.0,
            }
        }
        job_row = {
            "dedup_key": "acme|ds|remote",
            "title": "Data Scientist",
            "company": "Acme",
            "jd_full": "Looking for a data scientist.",
            "fit_analysis": None,
        }

        connect_calls = []
        original_connect = _sqlite3.connect

        def tracking_connect(path, **kwargs):
            connect_calls.append(path)
            return original_connect(path, **kwargs)

        with patch("job_finder.web.resume_generator.sqlite3.connect", side_effect=tracking_connect):
            with patch("job_finder.web.resume_generator.call_claude") as mock_call:
                mock_call.return_value = (sample_resume, 0.01)
                with patch("job_finder.web.resume_generator.cost_gate", return_value=True):
                    mock_client_factory = MagicMock()
                    mock_client_factory.return_value = MagicMock()
                    _generate_single_variant(
                        tmp_db_path,
                        mock_client_factory,
                        job_row,
                        sample_resume_data,
                        "impact_focused",
                        config,
                    )

        assert len(connect_calls) >= 1, (
            "_generate_single_variant must open its own sqlite3 connection"
        )
        assert tmp_db_path in connect_calls, (
            f"Expected connection to {tmp_db_path}, got connections to: {connect_calls}"
        )


class TestPartialFailure:
    """Partial failure: 1 variant fails, synthesis uses remaining 2."""

    def _make_sample_resume(self, suffix=""):
        return {
            "name": f"Jane Doe{suffix}",
            "contact_line": "jane@example.com",
            "summary": f"Experienced scientist{suffix}.",
            "skills": ["Python"],
            "positions": [],
            "education": [],
        }

    def test_partial_failure_still_synthesizes_remaining_variants(self, tmp_db_path, sample_resume_data):
        """When 1 of 3 variants fails, synthesis runs with the 2 remaining variants."""
        import sqlite3 as _sqlite3
        from unittest.mock import MagicMock, patch

        from job_finder.web.db_migrate import run_migrations
        from job_finder.web.resume_generator import STRATEGY_POOL, generate_resume_multi

        run_migrations(tmp_db_path)

        config = {
            "scoring": {
                "models": {
                    "haiku": "claude-haiku-4-5",
                    "sonnet": "claude-sonnet-4-6",
                },
                "monthly_budget_usd": 25.0,
            }
        }
        job_row = {
            "dedup_key": "acme|ds|remote",
            "title": "Data Scientist",
            "company": "Acme",
            "jd_full": "Looking for a data scientist.",
            "fit_analysis": None,
            "sonnet_score": 85.0,
        }

        sample_resume = self._make_sample_resume()
        call_count = {"n": 0}

        def mock_generate_variant(db_path, client_factory, job_row, profile, strategy, config):
            call_count["n"] += 1
            if call_count["n"] == 1:
                raise RuntimeError("Variant 1 failed")
            return self._make_sample_resume(suffix=f"_{strategy}")

        with patch("job_finder.web.resume_generator._generate_single_variant", side_effect=mock_generate_variant):
            with patch("job_finder.web.resume_generator._haiku_select_strategies") as mock_strat:
                mock_strat.return_value = STRATEGY_POOL[:3]
                with patch("job_finder.web.resume_generator._synthesize_variants") as mock_synth:
                    mock_synth.return_value = sample_resume
                    with patch("job_finder.web.resume_generator.anthropic") as mock_ant:
                        mock_ant.Anthropic.return_value = MagicMock()
                        result = generate_resume_multi(tmp_db_path, job_row, sample_resume_data, config)

        # _synthesize_variants should have been called with 2 variants (not 3)
        assert mock_synth.called, "_synthesize_variants must be called even with partial failure"
        synth_call_args = mock_synth.call_args
        # First positional arg after db_path is variants list
        variants_arg = synth_call_args[0][1] if synth_call_args[0] else synth_call_args.kwargs.get("variants")
        assert variants_arg is not None, "variants argument not passed to _synthesize_variants"
        assert len(variants_arg) == 2, (
            f"Expected 2 variants passed to synthesis (1 failed), got: {len(variants_arg)}"
        )
        assert result is not None

    def test_total_failure_raises_runtime_error(self, tmp_db_path, sample_resume_data):
        """When all 3 variants fail, generate_resume_multi raises RuntimeError."""
        from unittest.mock import MagicMock, patch

        from job_finder.web.db_migrate import run_migrations
        from job_finder.web.resume_generator import STRATEGY_POOL, generate_resume_multi

        run_migrations(tmp_db_path)

        config = {
            "scoring": {
                "models": {
                    "haiku": "claude-haiku-4-5",
                    "sonnet": "claude-sonnet-4-6",
                },
                "monthly_budget_usd": 25.0,
            }
        }
        job_row = {
            "dedup_key": "acme|ds|remote",
            "title": "Data Scientist",
            "company": "Acme",
            "jd_full": "Looking for a data scientist.",
            "fit_analysis": None,
            "sonnet_score": 85.0,
        }

        with patch("job_finder.web.resume_generator._generate_single_variant", side_effect=RuntimeError("all fail")):
            with patch("job_finder.web.resume_generator._haiku_select_strategies") as mock_strat:
                mock_strat.return_value = STRATEGY_POOL[:3]
                with patch("job_finder.web.resume_generator.anthropic") as mock_ant:
                    mock_ant.Anthropic.return_value = MagicMock()
                    with pytest.raises(RuntimeError, match="All resume variants failed"):
                        generate_resume_multi(tmp_db_path, job_row, sample_resume_data, config)


class TestSynthesisPass:
    """_synthesize_variants merges variants into a RESUME_SCHEMA-conforming dict."""

    def test_synthesize_variants_returns_resume_schema_dict(self, tmp_db_path):
        """_synthesize_variants returns a dict with all RESUME_SCHEMA required keys."""
        import sqlite3 as _sqlite3
        from unittest.mock import MagicMock, patch

        from job_finder.web.db_migrate import run_migrations
        from job_finder.web.resume_generator import RESUME_SCHEMA, _synthesize_variants

        run_migrations(tmp_db_path)

        config = {
            "scoring": {
                "models": {"sonnet": "claude-sonnet-4-6"},
                "monthly_budget_usd": 25.0,
            }
        }
        job_row = {
            "dedup_key": "acme|ds|remote",
            "title": "Data Scientist",
            "company": "Acme",
            "jd_full": "Looking for a data scientist.",
        }

        variants = [
            {
                "name": "Jane Doe",
                "contact_line": "jane@example.com",
                "summary": "Impact-focused summary.",
                "skills": ["Python", "SQL"],
                "positions": [],
                "education": [],
            },
            {
                "name": "Jane Doe",
                "contact_line": "jane@example.com",
                "summary": "Technical-depth summary.",
                "skills": ["ML", "Spark"],
                "positions": [],
                "education": [],
            },
        ]

        expected_result = {
            "name": "Jane Doe",
            "contact_line": "jane@example.com",
            "summary": "Best merged summary.",
            "skills": ["Python", "SQL", "ML"],
            "positions": [],
            "education": [],
        }

        with patch("job_finder.web.resume_generator.call_claude") as mock_call:
            mock_call.return_value = (expected_result, 0.05)
            with patch("job_finder.web.resume_generator.cost_gate", return_value=True):
                with patch("job_finder.web.resume_generator.anthropic") as mock_ant:
                    mock_ant.Anthropic.return_value = MagicMock()
                    result = _synthesize_variants(tmp_db_path, variants, job_row, config)

        required_keys = RESUME_SCHEMA.get("required", [])
        for key in required_keys:
            assert key in result, f"Synthesized resume missing required RESUME_SCHEMA key: {key}"

    def test_synthesize_variants_calls_call_claude_with_resume_synthesis_purpose(self, tmp_db_path):
        """_synthesize_variants calls call_claude with purpose='resume_synthesis'."""
        import sqlite3 as _sqlite3
        from unittest.mock import MagicMock, patch

        from job_finder.web.db_migrate import run_migrations
        from job_finder.web.resume_generator import _synthesize_variants

        run_migrations(tmp_db_path)

        config = {
            "scoring": {
                "models": {"sonnet": "claude-sonnet-4-6"},
                "monthly_budget_usd": 25.0,
            }
        }
        job_row = {
            "dedup_key": "acme|ds|remote",
            "title": "Data Scientist",
            "company": "Acme",
            "jd_full": "Looking for a data scientist.",
        }

        variants = [
            {
                "name": "Jane Doe",
                "contact_line": "",
                "summary": "test",
                "skills": [],
                "positions": [],
                "education": [],
            }
        ]

        synth_result = {
            "name": "Jane Doe",
            "contact_line": "",
            "summary": "merged",
            "skills": [],
            "positions": [],
            "education": [],
        }

        with patch("job_finder.web.resume_generator.call_claude") as mock_call:
            mock_call.return_value = (synth_result, 0.05)
            with patch("job_finder.web.resume_generator.cost_gate", return_value=True):
                with patch("job_finder.web.resume_generator.anthropic") as mock_ant:
                    mock_ant.Anthropic.return_value = MagicMock()
                    _synthesize_variants(tmp_db_path, variants, job_row, config)

        assert mock_call.called, "call_claude not called in _synthesize_variants"
        purpose = mock_call.call_args.kwargs.get("purpose") or mock_call.call_args[0][7]
        assert purpose == "resume_synthesis", (
            f"Expected purpose='resume_synthesis', got: {purpose}"
        )

    def test_synthesize_variants_passes_all_variants_to_call_claude(self, tmp_db_path):
        """_synthesize_variants includes all variant content in the user message."""
        import sqlite3 as _sqlite3
        from unittest.mock import MagicMock, patch

        from job_finder.web.db_migrate import run_migrations
        from job_finder.web.resume_generator import _synthesize_variants

        run_migrations(tmp_db_path)

        config = {
            "scoring": {
                "models": {"sonnet": "claude-sonnet-4-6"},
                "monthly_budget_usd": 25.0,
            }
        }
        job_row = {
            "dedup_key": "acme|ds|remote",
            "title": "Data Scientist",
            "company": "Acme",
            "jd_full": "Looking for a data scientist.",
        }

        # Two distinct variants with unique identifiers
        variants = [
            {
                "name": "Jane Doe",
                "contact_line": "",
                "summary": "UNIQUE_VARIANT_1_SUMMARY",
                "skills": [],
                "positions": [],
                "education": [],
            },
            {
                "name": "Jane Doe",
                "contact_line": "",
                "summary": "UNIQUE_VARIANT_2_SUMMARY",
                "skills": [],
                "positions": [],
                "education": [],
            },
        ]

        synth_result = {
            "name": "Jane Doe",
            "contact_line": "",
            "summary": "merged",
            "skills": [],
            "positions": [],
            "education": [],
        }

        captured_messages = {}

        def capturing_call_claude(**kwargs):
            captured_messages["messages"] = kwargs.get("messages", [])
            return (synth_result, 0.05)

        with patch("job_finder.web.resume_generator.call_claude", side_effect=capturing_call_claude):
            with patch("job_finder.web.resume_generator.cost_gate", return_value=True):
                with patch("job_finder.web.resume_generator.anthropic") as mock_ant:
                    mock_ant.Anthropic.return_value = MagicMock()
                    _synthesize_variants(tmp_db_path, variants, job_row, config)

        messages = captured_messages.get("messages", [])
        assert messages, "No messages passed to call_claude from _synthesize_variants"
        user_content = " ".join(m.get("content", "") for m in messages if m.get("role") == "user")
        assert "UNIQUE_VARIANT_1_SUMMARY" in user_content, (
            "Variant 1 content not included in synthesis prompt"
        )
        assert "UNIQUE_VARIANT_2_SUMMARY" in user_content, (
            "Variant 2 content not included in synthesis prompt"
        )


class TestScoreThresholdDispatch:
    """_generate_resume_background dispatches single vs multi based on sonnet_score."""

    def test_dispatches_multi_when_sonnet_score_above_threshold(self, tmp_db_path, sample_resume_data):
        """_generate_resume_background calls generate_resume_multi when sonnet_score >= 80."""
        import sqlite3 as _sqlite3
        from unittest.mock import MagicMock, patch

        from job_finder.web.db_migrate import run_migrations
        from job_finder.web.resume_generator import _generate_resume_background

        run_migrations(tmp_db_path)
        conn = _sqlite3.connect(tmp_db_path)
        conn.execute(
            "INSERT INTO resume_generations (job_id, generated_at, model, status, generation_type) "
            "VALUES (?, ?, ?, ?, ?)",
            ("acme|ds|remote", "2026-03-11T00:00:00", "claude-sonnet-4-6", "pending", "single"),
        )
        conn.commit()
        gen_id = conn.execute("SELECT last_insert_rowid()").fetchone()[0]
        conn.close()

        job_row = {
            "dedup_key": "acme|ds|remote",
            "title": "Data Scientist",
            "company": "Acme",
            "jd_full": "Looking for a data scientist.",
            "fit_analysis": None,
            "sonnet_score": 85.0,  # >= 80 threshold
        }
        config = {
            "scoring": {
                "models": {"sonnet": "claude-sonnet-4-6", "haiku": "claude-haiku-4-5"},
                "monthly_budget_usd": 25.0,
                "multi_version_threshold": 80,
            },
            "drive": {"folder_id": "test-folder", "convert_to_gdoc": True},
        }

        mock_resume = {
            "name": "Jane Doe",
            "contact_line": "",
            "summary": "test",
            "skills": [],
            "positions": [],
            "education": [],
        }

        with patch("job_finder.web.resume_generator.generate_resume_multi", return_value=mock_resume) as mock_multi:
            with patch("job_finder.web.resume_generator.generate_resume_single") as mock_single:
                with patch("job_finder.web.resume_generator.build_resume_docx") as mock_docx:
                    mock_docx.return_value = __import__("io").BytesIO(b"fake-docx")
                    with patch("job_finder.web.resume_generator.get_drive_service"):
                        with patch("job_finder.web.resume_generator.upload_to_drive", return_value="https://docs.google.com/xyz"):
                            _generate_resume_background(tmp_db_path, gen_id, job_row, sample_resume_data, config)

        assert mock_multi.called, "generate_resume_multi should be called for sonnet_score=85 >= threshold=80"
        assert not mock_single.called, "generate_resume_single should NOT be called for high-score job"

    def test_dispatches_single_when_sonnet_score_below_threshold(self, tmp_db_path, sample_resume_data):
        """_generate_resume_background calls generate_resume_single when sonnet_score < 80."""
        import sqlite3 as _sqlite3
        from unittest.mock import MagicMock, patch

        from job_finder.web.db_migrate import run_migrations
        from job_finder.web.resume_generator import _generate_resume_background

        run_migrations(tmp_db_path)
        conn = _sqlite3.connect(tmp_db_path)
        conn.execute(
            "INSERT INTO resume_generations (job_id, generated_at, model, status, generation_type) "
            "VALUES (?, ?, ?, ?, ?)",
            ("acme|ds|remote", "2026-03-11T00:00:00", "claude-sonnet-4-6", "pending", "single"),
        )
        conn.commit()
        gen_id = conn.execute("SELECT last_insert_rowid()").fetchone()[0]
        conn.close()

        job_row = {
            "dedup_key": "acme|ds|remote",
            "title": "Data Scientist",
            "company": "Acme",
            "jd_full": "Looking for a data scientist.",
            "fit_analysis": None,
            "sonnet_score": 70.0,  # < 80 threshold
        }
        config = {
            "scoring": {
                "models": {"sonnet": "claude-sonnet-4-6", "haiku": "claude-haiku-4-5"},
                "monthly_budget_usd": 25.0,
                "multi_version_threshold": 80,
            },
            "drive": {"folder_id": "test-folder", "convert_to_gdoc": True},
        }

        mock_resume = {
            "name": "Jane Doe",
            "contact_line": "",
            "summary": "test",
            "skills": [],
            "positions": [],
            "education": [],
        }

        with patch("job_finder.web.resume_generator.generate_resume_multi") as mock_multi:
            with patch("job_finder.web.resume_generator.generate_resume_single", return_value=mock_resume) as mock_single:
                with patch("job_finder.web.resume_generator.build_resume_docx") as mock_docx:
                    mock_docx.return_value = __import__("io").BytesIO(b"fake-docx")
                    with patch("job_finder.web.resume_generator.get_drive_service"):
                        with patch("job_finder.web.resume_generator.upload_to_drive", return_value="https://docs.google.com/xyz"):
                            _generate_resume_background(tmp_db_path, gen_id, job_row, sample_resume_data, config)

        assert mock_single.called, "generate_resume_single should be called for sonnet_score=70 < threshold=80"
        assert not mock_multi.called, "generate_resume_multi should NOT be called for low-score job"

    def test_generation_type_set_to_multi_when_above_threshold(self, tmp_db_path, sample_resume_data):
        """_generate_resume_background sets generation_type='multi' in DB for high-score job."""
        import sqlite3 as _sqlite3
        from unittest.mock import MagicMock, patch

        from job_finder.web.db_migrate import run_migrations
        from job_finder.web.resume_generator import _generate_resume_background

        run_migrations(tmp_db_path)
        conn = _sqlite3.connect(tmp_db_path)
        conn.execute(
            "INSERT INTO resume_generations (job_id, generated_at, model, status, generation_type) "
            "VALUES (?, ?, ?, ?, ?)",
            ("acme|ds|remote", "2026-03-11T00:00:00", "claude-sonnet-4-6", "pending", "single"),
        )
        conn.commit()
        gen_id = conn.execute("SELECT last_insert_rowid()").fetchone()[0]
        conn.close()

        job_row = {
            "dedup_key": "acme|ds|remote",
            "title": "Data Scientist",
            "company": "Acme",
            "jd_full": "Looking for a data scientist.",
            "fit_analysis": None,
            "sonnet_score": 90.0,
        }
        config = {
            "scoring": {
                "models": {"sonnet": "claude-sonnet-4-6", "haiku": "claude-haiku-4-5"},
                "monthly_budget_usd": 25.0,
                "multi_version_threshold": 80,
            },
            "drive": {"folder_id": "test-folder", "convert_to_gdoc": True},
        }

        mock_resume = {
            "name": "Jane Doe",
            "contact_line": "",
            "summary": "test",
            "skills": [],
            "positions": [],
            "education": [],
        }

        with patch("job_finder.web.resume_generator.generate_resume_multi", return_value=mock_resume):
            with patch("job_finder.web.resume_generator.build_resume_docx") as mock_docx:
                mock_docx.return_value = __import__("io").BytesIO(b"fake-docx")
                with patch("job_finder.web.resume_generator.get_drive_service"):
                    with patch("job_finder.web.resume_generator.upload_to_drive", return_value="https://docs.google.com/xyz"):
                        _generate_resume_background(tmp_db_path, gen_id, job_row, sample_resume_data, config)

        verify_conn = _sqlite3.connect(tmp_db_path)
        row = verify_conn.execute(
            "SELECT generation_type FROM resume_generations WHERE id = ?", (gen_id,)
        ).fetchone()
        verify_conn.close()

        assert row[0] == "multi", f"Expected generation_type='multi', got: {row[0]}"


# =============================================================================
# Plan 04 Tests: Quick Apply route
# =============================================================================

class TestQuickApply:
    """POST /jobs/<key>/quick-apply -- generate resume, open tabs, set status applied."""

    @pytest.fixture
    def app_with_existing_resume(self, tmp_db_path):
        """Create test app with a job that already has a done resume."""
        import sqlite3 as _sqlite3

        from job_finder.web.db_migrate import run_migrations
        from job_finder.web import create_app

        run_migrations(tmp_db_path)

        conn = _sqlite3.connect(tmp_db_path)
        conn.execute(
            """INSERT INTO jobs
                (dedup_key, title, company, location, sources, source_urls,
                 source_id, first_seen, last_seen, score, score_breakdown,
                 user_interest, pipeline_status, sonnet_score, jd_full)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
            (
                "acme|senior-ds|remote",
                "Senior Data Scientist",
                "Acme Corp",
                "Remote",
                '["linkedin"]',
                '["https://linkedin.com/jobs/1234"]',
                "1234",
                "2026-03-01T00:00:00",
                "2026-03-11T00:00:00",
                8.5,
                "{}",
                "reviewing",
                "reviewing",
                85.0,
                "Full job description text here.",
            ),
        )
        conn.commit()
        # Insert an existing done resume
        conn.execute(
            "INSERT INTO resume_generations (job_id, generated_at, model, status, doc_url, generation_type) "
            "VALUES (?, ?, ?, ?, ?, ?)",
            (
                "acme|senior-ds|remote",
                "2026-03-10T00:00:00Z",
                "claude-sonnet-4-6",
                "done",
                "https://docs.google.com/doc/existing123",
                "single",
            ),
        )
        conn.commit()
        conn.close()

        cfg = {
            "db": {"path": tmp_db_path},
            "scoring": {"models": {"sonnet": "claude-sonnet-4-6"}, "monthly_budget_usd": 25.0},
            "drive": {"folder_id": "test-folder", "convert_to_gdoc": True},
        }
        app = create_app(config=cfg)
        app.config["TESTING"] = True
        return app

    def test_quick_apply_with_existing_resume_returns_200(self, app_with_existing_resume):
        """POST /quick-apply with existing done resume returns 200."""
        from urllib.parse import quote

        dedup_key = "acme|senior-ds|remote"
        app = app_with_existing_resume

        with app.test_client() as client:
            resp = client.post(f"/jobs/{quote(dedup_key, safe='')}/quick-apply")

        assert resp.status_code == 200, f"Expected 200, got: {resp.status_code}"

    def test_quick_apply_with_existing_resume_contains_doc_url(self, app_with_existing_resume):
        """POST /quick-apply response contains the existing doc_url."""
        from urllib.parse import quote

        dedup_key = "acme|senior-ds|remote"
        app = app_with_existing_resume

        with app.test_client() as client:
            resp = client.post(f"/jobs/{quote(dedup_key, safe='')}/quick-apply")

        body = resp.data.decode()
        assert "docs.google.com/doc/existing123" in body, (
            "Response should contain the existing doc URL"
        )

    def test_quick_apply_with_existing_resume_sets_status_applied(self, app_with_existing_resume, tmp_db_path):
        """POST /quick-apply sets pipeline_status to 'applied'."""
        import sqlite3 as _sqlite3
        from urllib.parse import quote

        dedup_key = "acme|senior-ds|remote"
        app = app_with_existing_resume

        with app.test_client() as client:
            client.post(f"/jobs/{quote(dedup_key, safe='')}/quick-apply")

        conn = _sqlite3.connect(tmp_db_path)
        row = conn.execute(
            "SELECT pipeline_status FROM jobs WHERE dedup_key = ?", (dedup_key,)
        ).fetchone()
        conn.close()

        assert row[0] == "applied", f"Expected pipeline_status='applied', got: {row[0]}"

    def test_quick_apply_without_existing_resume_triggers_generation(self, tmp_db_path):
        """POST /quick-apply without existing resume triggers generation (mocked)."""
        import sqlite3 as _sqlite3
        from unittest.mock import patch
        from urllib.parse import quote

        from job_finder.web.db_migrate import run_migrations
        from job_finder.web import create_app

        run_migrations(tmp_db_path)
        conn = _sqlite3.connect(tmp_db_path)
        conn.execute(
            """INSERT INTO jobs
                (dedup_key, title, company, location, sources, source_urls,
                 source_id, first_seen, last_seen, score, score_breakdown,
                 user_interest, pipeline_status, sonnet_score, jd_full)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
            (
                "acme|senior-ds|remote",
                "Senior Data Scientist",
                "Acme Corp",
                "Remote",
                '["linkedin"]',
                '["https://linkedin.com/jobs/1234"]',
                "1234",
                "2026-03-01T00:00:00",
                "2026-03-11T00:00:00",
                8.5,
                "{}",
                "reviewing",
                "reviewing",
                85.0,
                "Full job description text here.",
            ),
        )
        conn.commit()
        conn.close()

        cfg = {
            "db": {"path": tmp_db_path},
            "scoring": {"models": {"sonnet": "claude-sonnet-4-6"}, "monthly_budget_usd": 25.0},
            "drive": {"folder_id": "test-folder", "convert_to_gdoc": True},
        }
        app = create_app(config=cfg)
        app.config["TESTING"] = True

        mock_resume = {
            "name": "Jane Doe",
            "contact_line": "jane@example.com",
            "summary": "Experienced data scientist.",
            "skills": ["Python"],
            "positions": [],
            "education": [],
        }

        with patch("job_finder.web.blueprints.resume.generate_resume_single", return_value=mock_resume):
            with patch("job_finder.web.blueprints.resume.build_resume_docx") as mock_docx:
                mock_docx.return_value = __import__("io").BytesIO(b"fake-docx")
                with patch("job_finder.web.blueprints.resume.get_drive_service"):
                    with patch("job_finder.web.blueprints.resume.upload_to_drive", return_value="https://docs.google.com/doc/new456"):
                        with patch("job_finder.web.blueprints.resume.load_profile", return_value={}):
                            with patch("job_finder.web.blueprints.resume.anthropic.Anthropic"):
                                with app.test_client() as client:
                                    resp = client.post(f"/jobs/{quote('acme|senior-ds|remote', safe='')}/quick-apply")

        assert resp.status_code == 200, f"Expected 200, got: {resp.status_code}"
        body = resp.data.decode()
        assert "Applied" in body or "docs.google.com" in body, (
            "Response should show applied confirmation"
        )

    def test_quick_apply_returns_400_when_no_sonnet_score(self, tmp_db_path):
        """POST /quick-apply returns 400 when job has no sonnet_score."""
        import sqlite3 as _sqlite3
        from urllib.parse import quote

        from job_finder.web.db_migrate import run_migrations
        from job_finder.web import create_app

        run_migrations(tmp_db_path)

        conn = _sqlite3.connect(tmp_db_path)
        conn.execute(
            """INSERT INTO jobs
                (dedup_key, title, company, location, sources, source_urls,
                 source_id, first_seen, last_seen, score, score_breakdown,
                 user_interest, pipeline_status, haiku_score)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
            (
                "haiku|only|job",
                "Data Analyst",
                "Some Corp",
                "Remote",
                '["linkedin"]',
                '["https://linkedin.com/jobs/5678"]',
                "5678",
                "2026-03-01T00:00:00",
                "2026-03-11T00:00:00",
                6.0,
                "{}",
                "reviewing",
                "reviewing",
                65.0,
            ),
        )
        conn.commit()
        conn.close()

        cfg = {
            "db": {"path": tmp_db_path},
            "scoring": {"models": {"sonnet": "claude-sonnet-4-6"}, "monthly_budget_usd": 25.0},
        }
        app = create_app(config=cfg)
        app.config["TESTING"] = True

        with app.test_client() as client:
            resp = client.post(f"/jobs/{quote('haiku|only|job', safe='')}/quick-apply")

        assert resp.status_code == 400, f"Expected 400, got: {resp.status_code}"

    def test_quick_apply_google_search_fallback_when_no_source_urls(self, tmp_db_path):
        """POST /quick-apply uses Google search fallback when source_urls is empty."""
        import sqlite3 as _sqlite3
        from urllib.parse import quote

        from job_finder.web.db_migrate import run_migrations
        from job_finder.web import create_app

        run_migrations(tmp_db_path)

        conn = _sqlite3.connect(tmp_db_path)
        conn.execute(
            """INSERT INTO jobs
                (dedup_key, title, company, location, sources, source_urls,
                 source_id, first_seen, last_seen, score, score_breakdown,
                 user_interest, pipeline_status, sonnet_score, jd_full)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
            (
                "nourls|job|remote",
                "Staff Analyst",
                "NoURLs Corp",
                "Remote",
                '[]',
                '[]',
                "",
                "2026-03-01T00:00:00",
                "2026-03-11T00:00:00",
                8.0,
                "{}",
                "reviewing",
                "reviewing",
                80.0,
                "Job description here.",
            ),
        )
        conn.commit()
        # Insert existing done resume to skip generation
        conn.execute(
            "INSERT INTO resume_generations (job_id, generated_at, model, status, doc_url, generation_type) "
            "VALUES (?, ?, ?, ?, ?, ?)",
            (
                "nourls|job|remote",
                "2026-03-10T00:00:00Z",
                "claude-sonnet-4-6",
                "done",
                "https://docs.google.com/doc/fallback-test",
                "single",
            ),
        )
        conn.commit()
        conn.close()

        cfg = {
            "db": {"path": tmp_db_path},
            "scoring": {"models": {"sonnet": "claude-sonnet-4-6"}, "monthly_budget_usd": 25.0},
        }
        app = create_app(config=cfg)
        app.config["TESTING"] = True

        with app.test_client() as client:
            resp = client.post(f"/jobs/{quote('nourls|job|remote', safe='')}/quick-apply")

        body = resp.data.decode()
        assert "google.com/search" in body, (
            "Response should contain Google search fallback URL"
        )


class TestQuickApplyResponse:
    """_quick_apply_response.html contains expected JS tab-opening and confirmation links."""

    def test_response_contains_window_open_for_doc_url(self, tmp_db_path):
        """Quick apply response HTML contains window.open for doc_url."""
        import sqlite3 as _sqlite3
        from urllib.parse import quote

        from job_finder.web.db_migrate import run_migrations
        from job_finder.web import create_app

        run_migrations(tmp_db_path)

        conn = _sqlite3.connect(tmp_db_path)
        conn.execute(
            """INSERT INTO jobs
                (dedup_key, title, company, location, sources, source_urls,
                 source_id, first_seen, last_seen, score, score_breakdown,
                 user_interest, pipeline_status, sonnet_score, jd_full)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
            (
                "acme|senior-ds|remote",
                "Senior Data Scientist",
                "Acme Corp",
                "Remote",
                '["linkedin"]',
                '["https://linkedin.com/jobs/1234"]',
                "1234",
                "2026-03-01T00:00:00",
                "2026-03-11T00:00:00",
                8.5,
                "{}",
                "reviewing",
                "reviewing",
                85.0,
                "Full job description.",
            ),
        )
        conn.commit()
        conn.execute(
            "INSERT INTO resume_generations (job_id, generated_at, model, status, doc_url, generation_type) "
            "VALUES (?, ?, ?, ?, ?, ?)",
            (
                "acme|senior-ds|remote",
                "2026-03-10T00:00:00Z",
                "claude-sonnet-4-6",
                "done",
                "https://docs.google.com/doc/abc123",
                "single",
            ),
        )
        conn.commit()
        conn.close()

        cfg = {
            "db": {"path": tmp_db_path},
            "scoring": {"models": {"sonnet": "claude-sonnet-4-6"}, "monthly_budget_usd": 25.0},
        }
        app = create_app(config=cfg)
        app.config["TESTING"] = True

        with app.test_client() as client:
            resp = client.post(f"/jobs/{quote('acme|senior-ds|remote', safe='')}/quick-apply")

        body = resp.data.decode()
        assert "window.open" in body, "Response must contain window.open for tab opening"
        assert "docs.google.com/doc/abc123" in body, "Response must contain doc_url in window.open"

    def test_response_contains_window_open_for_app_url(self, tmp_db_path):
        """Quick apply response HTML contains window.open for app_url."""
        import sqlite3 as _sqlite3
        from urllib.parse import quote

        from job_finder.web.db_migrate import run_migrations
        from job_finder.web import create_app

        run_migrations(tmp_db_path)

        conn = _sqlite3.connect(tmp_db_path)
        conn.execute(
            """INSERT INTO jobs
                (dedup_key, title, company, location, sources, source_urls,
                 source_id, first_seen, last_seen, score, score_breakdown,
                 user_interest, pipeline_status, sonnet_score, jd_full)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
            (
                "acme|senior-ds|remote",
                "Senior Data Scientist",
                "Acme Corp",
                "Remote",
                '["linkedin"]',
                '["https://linkedin.com/jobs/9999"]',
                "9999",
                "2026-03-01T00:00:00",
                "2026-03-11T00:00:00",
                8.5,
                "{}",
                "reviewing",
                "reviewing",
                85.0,
                "Full job description.",
            ),
        )
        conn.commit()
        conn.execute(
            "INSERT INTO resume_generations (job_id, generated_at, model, status, doc_url, generation_type) "
            "VALUES (?, ?, ?, ?, ?, ?)",
            (
                "acme|senior-ds|remote",
                "2026-03-10T00:00:00Z",
                "claude-sonnet-4-6",
                "done",
                "https://docs.google.com/doc/abc123",
                "single",
            ),
        )
        conn.commit()
        conn.close()

        cfg = {
            "db": {"path": tmp_db_path},
            "scoring": {"models": {"sonnet": "claude-sonnet-4-6"}, "monthly_budget_usd": 25.0},
        }
        app = create_app(config=cfg)
        app.config["TESTING"] = True

        with app.test_client() as client:
            resp = client.post(f"/jobs/{quote('acme|senior-ds|remote', safe='')}/quick-apply")

        body = resp.data.decode()
        assert "window.open" in body, "Response must contain window.open for tab opening"
        assert "linkedin.com/jobs/9999" in body, "Response must contain app_url in window.open"


# =============================================================================
# Phase 08 Plan 01 Tests: Interview prep trigger helper and Quick Apply wiring
# =============================================================================

class TestInterviewPrepTrigger:
    """Tests for trigger_interview_prep_if_applied shared helper and Quick Apply wiring."""

    def test_helper_spawns_thread_when_applied_and_not_testing(self):
        """trigger_interview_prep_if_applied spawns a daemon thread when new_status='applied' and testing=False."""
        from unittest.mock import patch, MagicMock
        from job_finder.web.blueprints import trigger_interview_prep_if_applied

        with patch("job_finder.web.blueprints.threading") as mock_threading:
            mock_thread = MagicMock()
            mock_threading.Thread.return_value = mock_thread

            trigger_interview_prep_if_applied(
                dedup_key="acme|ds|remote",
                new_status="applied",
                db_path="/tmp/test.db",
                config={},
                testing=False,
            )

        mock_threading.Thread.assert_called_once()
        mock_thread.start.assert_called_once()

    def test_helper_is_noop_when_status_not_applied(self):
        """trigger_interview_prep_if_applied does not spawn thread when new_status != 'applied'."""
        from unittest.mock import patch
        from job_finder.web.blueprints import trigger_interview_prep_if_applied

        with patch("job_finder.web.blueprints.threading") as mock_threading:
            trigger_interview_prep_if_applied(
                dedup_key="acme|ds|remote",
                new_status="reviewing",
                db_path="/tmp/test.db",
                config={},
                testing=False,
            )

        mock_threading.Thread.assert_not_called()

    def test_helper_is_noop_when_testing_is_true(self):
        """trigger_interview_prep_if_applied does not spawn thread when testing=True."""
        from unittest.mock import patch
        from job_finder.web.blueprints import trigger_interview_prep_if_applied

        with patch("job_finder.web.blueprints.threading") as mock_threading:
            trigger_interview_prep_if_applied(
                dedup_key="acme|ds|remote",
                new_status="applied",
                db_path="/tmp/test.db",
                config={},
                testing=True,
            )

        mock_threading.Thread.assert_not_called()

    def test_quick_apply_calls_trigger_helper(self, tmp_db_path):
        """POST /quick-apply calls trigger_interview_prep_if_applied after setting status."""
        import sqlite3 as _sqlite3
        from unittest.mock import patch
        from urllib.parse import quote
        from job_finder.web import create_app
        from job_finder.web.db_migrate import run_migrations

        run_migrations(tmp_db_path)

        conn = _sqlite3.connect(tmp_db_path)
        conn.execute(
            """INSERT INTO jobs
                (dedup_key, title, company, location, sources, source_urls,
                 source_id, first_seen, last_seen, score, score_breakdown,
                 user_interest, pipeline_status, sonnet_score, jd_full)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
            (
                "acme|senior-ds|remote",
                "Senior Data Scientist",
                "Acme Corp",
                "Remote",
                '["linkedin"]',
                '["https://linkedin.com/jobs/1234"]',
                "1234",
                "2026-03-01T00:00:00",
                "2026-03-11T00:00:00",
                8.5, "{}", "reviewing", "reviewing",
                85.0, "Full job description text here.",
            ),
        )
        conn.execute(
            "INSERT INTO resume_generations (job_id, generated_at, model, status, doc_url, generation_type) "
            "VALUES (?, ?, ?, ?, ?, ?)",
            (
                "acme|senior-ds|remote", "2026-03-10T00:00:00Z",
                "claude-sonnet-4-6", "done",
                "https://docs.google.com/doc/existing123", "single",
            ),
        )
        conn.commit()
        conn.close()

        cfg = {
            "db": {"path": tmp_db_path},
            "scoring": {"models": {"sonnet": "claude-sonnet-4-6"}, "monthly_budget_usd": 25.0},
            "drive": {"folder_id": "test-folder", "convert_to_gdoc": True},
        }
        app = create_app(config=cfg)
        app.config["TESTING"] = True

        with patch(
            "job_finder.web.blueprints.resume.trigger_interview_prep_if_applied"
        ) as mock_trigger:
            with app.test_client() as client:
                resp = client.post(f"/jobs/{quote('acme|senior-ds|remote', safe='')}/quick-apply")

        assert resp.status_code == 200, f"Expected 200, got: {resp.status_code}"
        mock_trigger.assert_called_once()
        call_kwargs = mock_trigger.call_args
        # Verify dedup_key and new_status="applied" were passed
        args, kwargs = call_kwargs
        assert args[0] == "acme|senior-ds|remote" or kwargs.get("dedup_key") == "acme|senior-ds|remote"


# =============================================================================
# Plan 04 Tests: Settings Resume & Drive section
# =============================================================================

class TestSettingsResumeFormat:
    """Settings page: Resume & Drive section saves and displays correctly."""

    @pytest.fixture
    def settings_app(self, tmp_db_path, tmp_path):
        """Create test app with a real config.yaml that settings can read/write."""
        import yaml
        from job_finder.web import create_app

        config_path = str(tmp_path / "config.yaml")
        initial_config = {
            "profile": {
                "target_titles": ["Data Scientist"],
                "target_locations": ["Remote"],
                "min_salary": 150000,
                "industries": [],
                "exclusions": {"title_keywords": [], "companies": []},
                "skills": [],
            },
            "sources": {
                "gmail": {
                    "enabled": False,
                    "lookback_days": 7,
                    "senders": {
                        "linkedin_alerts": "",
                        "linkedin_jobs": "",
                        "glassdoor": "",
                        "indeed": "",
                        "ziprecruiter": "",
                    },
                },
                "serpapi": {"enabled": False, "api_key": "", "queries": []},
                "jsearch": {"enabled": False, "rapidapi_key": ""},
            },
            "scoring": {
                "weights": {
                    "title_match": 0.30,
                    "seniority_alignment": 0.20,
                    "location_fit": 0.15,
                    "salary_range": 0.15,
                    "industry_relevance": 0.10,
                    "company_signals": 0.05,
                    "recency": 0.05,
                },
                "min_score_threshold": 40,
                "monthly_budget_usd": 25.0,
                "haiku_threshold": 55,
                "models": {"haiku": "claude-haiku-4-5", "sonnet": "claude-sonnet-4-6"},
                "multi_version_threshold": 80,
            },
            "output": {"default_format": "cli", "markdown_path": "reports/", "max_results": 50},
            "db": {"path": "jobs.db"},
            "drive": {"folder_id": "old-folder-id", "convert_to_gdoc": True},
        }
        with open(config_path, "w") as f:
            yaml.dump(initial_config, f)

        cfg = {**initial_config, "db": {"path": tmp_db_path}}
        app = create_app(config=cfg)
        app.config["TESTING"] = True

        # Override config path to point to our temp file
        import job_finder.web.blueprints.settings as settings_mod
        original_path = settings_mod._CONFIG_PATH
        settings_mod._CONFIG_PATH = config_path

        yield app, config_path

        # Restore
        settings_mod._CONFIG_PATH = original_path

    def _base_form_data(self, **overrides):
        """Return a complete form data dict for settings save."""
        base = {
            "target_titles": "Data Scientist",
            "target_locations": "Remote",
            "min_salary": "150000",
            "industries": "",
            "exclusion_title_keywords": "",
            "exclusion_companies": "",
            "profile_skills": "",
            "gmail_enabled": "",
            "gmail_lookback_days": "7",
            "gmail_sender_linkedin_alerts": "",
            "gmail_sender_linkedin_jobs": "",
            "gmail_sender_glassdoor": "",
            "gmail_sender_indeed": "",
            "gmail_sender_ziprecruiter": "",
            "serpapi_enabled": "",
            "serpapi_api_key": "",
            "_serpapi_queries_present": "1",
            "jsearch_enabled": "",
            "jsearch_rapidapi_key": "",
            "weight_title_match": "0.30",
            "weight_seniority_alignment": "0.20",
            "weight_location_fit": "0.15",
            "weight_salary_range": "0.15",
            "weight_industry_relevance": "0.10",
            "weight_company_signals": "0.05",
            "weight_recency": "0.05",
            "min_score_threshold": "40",
            "monthly_budget_usd": "25.0",
            "haiku_threshold": "55",
            "model_haiku": "claude-haiku-4-5",
            "model_sonnet": "claude-sonnet-4-6",
            "output_default_format": "cli",
            "output_markdown_path": "reports/",
            "output_max_results": "50",
            "db_path": "jobs.db",
            "drive_folder_id": "folder-id",
            "drive_convert_to_gdoc": "on",
            "multi_version_threshold": "80",
        }
        base.update(overrides)
        return base

    def test_post_settings_updates_drive_folder_id(self, settings_app):
        """POST /settings/save with drive_folder_id updates config.yaml."""
        import yaml

        app, config_path = settings_app

        form_data = self._base_form_data(drive_folder_id="my-new-folder-id")
        with app.test_client() as client:
            resp = client.post("/settings/save", data=form_data, follow_redirects=False)

        assert resp.status_code in (302, 200), f"Expected redirect, got: {resp.status_code}"

        with open(config_path) as f:
            saved = yaml.safe_load(f)

        assert saved.get("drive", {}).get("folder_id") == "my-new-folder-id", (
            f"drive.folder_id not updated, got: {saved.get('drive', {})}"
        )

    def test_post_settings_updates_convert_to_gdoc_false(self, settings_app):
        """POST /settings/save with drive_convert_to_gdoc=false sets False."""
        import yaml

        app, config_path = settings_app

        # Submit the "false" radio value (as the real form's .docx radio does)
        form_data = self._base_form_data(drive_convert_to_gdoc="false")
        with app.test_client() as client:
            resp = client.post("/settings/save", data=form_data, follow_redirects=False)

        assert resp.status_code in (302, 200), f"Expected redirect, got: {resp.status_code}"

        with open(config_path) as f:
            saved = yaml.safe_load(f)

        assert saved.get("drive", {}).get("convert_to_gdoc") is False, (
            f"drive.convert_to_gdoc should be False when unchecked, got: {saved.get('drive', {})}"
        )

    def test_get_settings_shows_drive_folder_id(self, settings_app):
        """GET /settings displays current drive.folder_id in the form."""
        app, _ = settings_app

        with app.test_client() as client:
            resp = client.get("/settings/")

        assert resp.status_code == 200, f"Expected 200, got: {resp.status_code}"
        body = resp.data.decode()
        assert "drive_folder_id" in body, "Settings page should have drive_folder_id input"
        assert "old-folder-id" in body, "Settings page should show current folder_id value"

    def test_get_settings_shows_resume_drive_section(self, settings_app):
        """GET /settings page contains Resume & Drive section."""
        app, _ = settings_app

        with app.test_client() as client:
            resp = client.get("/settings/")

        assert resp.status_code == 200
        body = resp.data.decode()
        assert "Resume" in body and "Drive" in body, (
            "Settings page must have Resume & Drive section"
        )


# =============================================================================
# Task 1 (Phase 08-02) Tests: Preference injection into resume generation paths
# =============================================================================

class TestPreferenceInjection:
    """Accepted resume preferences from Drive feedback loop injected into Sonnet prompts."""

    def _make_db_with_preferences(self, path: str, preferences: list) -> None:
        """Create a migrated DB and populate resume_preferences_detected table."""
        import sqlite3 as _sqlite3

        from job_finder.web.db_migrate import run_migrations

        run_migrations(path)
        conn = _sqlite3.connect(path)
        conn.execute(
            """CREATE TABLE IF NOT EXISTS resume_preferences_detected (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                job_id TEXT,
                preference_type TEXT,
                preference_text TEXT,
                example_before TEXT,
                example_after TEXT,
                accepted INTEGER DEFAULT 1,
                detected_at TEXT,
                applied_at TEXT DEFAULT NULL
            )"""
        )
        for pref in preferences:
            conn.execute(
                """INSERT INTO resume_preferences_detected
                   (job_id, preference_type, preference_text, accepted, detected_at, applied_at)
                   VALUES (?, ?, ?, ?, ?, ?)""",
                (
                    pref.get("job_id", "test-job"),
                    pref.get("preference_type", "phrasing"),
                    pref["preference_text"],
                    pref.get("accepted", 1),
                    pref.get("detected_at", "2026-01-01T00:00:00"),
                    pref.get("applied_at", None),
                ),
            )
        conn.commit()
        conn.close()

    def _make_job_row(self) -> dict:
        return {
            "dedup_key": "acme|senior-ds|remote",
            "title": "Senior Data Scientist",
            "company": "Acme Corp",
            "jd_full": "Looking for a data scientist with Python and ML experience.",
            "fit_analysis": None,
            "sonnet_score": 75.0,
        }

    def _make_profile(self) -> dict:
        return {
            "skills": ["Python", "SQL", "Machine Learning"],
            "positions": [
                {
                    "title": "Data Scientist",
                    "company": "Beta Inc",
                    "start_date": "2020",
                    "end_date": "Present",
                    "achievements": ["Built ML models"],
                    "skills": ["Python"],
                }
            ],
            "resume_preferences": {},
        }

    def _make_config(self) -> dict:
        return {
            "scoring": {
                "models": {"sonnet": "claude-sonnet-4-6"},
                "monthly_budget_usd": 25.0,
            }
        }

    def test_single_pass_includes_preferences_when_accepted_prefs_exist(self, tmp_db_path):
        """generate_resume_single appends 'Candidate Writing Preferences' when accepted prefs exist."""
        import sqlite3 as _sqlite3
        from unittest.mock import MagicMock, patch

        from job_finder.web.resume_generator import generate_resume_single

        self._make_db_with_preferences(tmp_db_path, [
            {"preference_text": "Use active voice in bullet points", "preference_type": "phrasing"},
            {"preference_text": "Start bullets with strong action verbs", "preference_type": "phrasing"},
        ])

        conn = _sqlite3.connect(tmp_db_path)
        conn.row_factory = _sqlite3.Row

        captured_messages = []

        def capturing_call_claude(**kwargs):
            captured_messages.extend(kwargs.get("messages", []))
            return (
                {
                    "name": "Jane Doe",
                    "contact_line": "jane@example.com",
                    "summary": "Experienced DS.",
                    "skills": ["Python"],
                    "positions": [],
                    "education": [],
                },
                0.01,
            )

        with patch("job_finder.web.resume_generator.call_claude", side_effect=capturing_call_claude):
            with patch("job_finder.web.resume_generator.cost_gate", return_value=True):
                mock_client = MagicMock()
                generate_resume_single(
                    mock_client,
                    self._make_job_row(),
                    self._make_profile(),
                    conn,
                    self._make_config(),
                )

        conn.close()

        assert captured_messages, "call_claude was not called"
        user_message = captured_messages[0]["content"]
        assert "Formatting Preferences" in user_message, (
            "Expected '## Formatting Preferences' section in user_message"
        )
        assert "Use active voice in bullet points" in user_message, (
            "Expected first preference text in user_message"
        )
        assert "Start bullets with strong action verbs" in user_message, (
            "Expected second preference text in user_message"
        )

    def test_single_pass_preferences_formatted_as_bullet_list_with_soft_guidelines(self, tmp_db_path):
        """Preferences are formatted as bullet list with 'soft guidelines' framing."""
        import sqlite3 as _sqlite3
        from unittest.mock import MagicMock, patch

        from job_finder.web.resume_generator import generate_resume_single

        self._make_db_with_preferences(tmp_db_path, [
            {"preference_text": "Quantify all achievements with numbers", "preference_type": "content_addition"},
        ])

        conn = _sqlite3.connect(tmp_db_path)
        conn.row_factory = _sqlite3.Row

        captured_messages = []

        def capturing_call_claude(**kwargs):
            captured_messages.extend(kwargs.get("messages", []))
            return (
                {
                    "name": "Jane Doe",
                    "contact_line": "",
                    "summary": "DS.",
                    "skills": [],
                    "positions": [],
                    "education": [],
                },
                0.01,
            )

        with patch("job_finder.web.resume_generator.call_claude", side_effect=capturing_call_claude):
            with patch("job_finder.web.resume_generator.cost_gate", return_value=True):
                mock_client = MagicMock()
                generate_resume_single(
                    mock_client,
                    self._make_job_row(),
                    self._make_profile(),
                    conn,
                    self._make_config(),
                )

        conn.close()

        assert captured_messages, "call_claude was not called"
        user_message = captured_messages[0]["content"]
        assert "soft guidelines" in user_message, (
            "Expected 'soft guidelines' framing in preferences section"
        )
        assert "- Quantify all achievements with numbers" in user_message, (
            "Expected preference formatted as bullet item '- <text>'"
        )

    def test_single_pass_no_preferences_section_when_table_empty(self, tmp_db_path):
        """generate_resume_single does NOT include preferences section when table is empty."""
        import sqlite3 as _sqlite3
        from unittest.mock import MagicMock, patch

        from job_finder.web.resume_generator import generate_resume_single

        self._make_db_with_preferences(tmp_db_path, [])

        conn = _sqlite3.connect(tmp_db_path)
        conn.row_factory = _sqlite3.Row

        captured_messages = []

        def capturing_call_claude(**kwargs):
            captured_messages.extend(kwargs.get("messages", []))
            return (
                {
                    "name": "Jane Doe",
                    "contact_line": "",
                    "summary": "DS.",
                    "skills": [],
                    "positions": [],
                    "education": [],
                },
                0.01,
            )

        with patch("job_finder.web.resume_generator.call_claude", side_effect=capturing_call_claude):
            with patch("job_finder.web.resume_generator.cost_gate", return_value=True):
                with patch("job_finder.web.resume_style_guide.load_style_guide", return_value={}):
                    mock_client = MagicMock()
                    generate_resume_single(
                        mock_client,
                        self._make_job_row(),
                        self._make_profile(),
                        conn,
                        self._make_config(),
                    )

        conn.close()

        assert captured_messages, "call_claude was not called"
        user_message = captured_messages[0]["content"]
        assert "Formatting Preferences" not in user_message, (
            "Should NOT include preferences section when table is empty"
        )

    def test_single_pass_excludes_preferences_with_applied_at_set(self, tmp_db_path):
        """Preferences with applied_at set are excluded from the prompt."""
        import sqlite3 as _sqlite3
        from unittest.mock import MagicMock, patch

        from job_finder.web.resume_generator import generate_resume_single

        self._make_db_with_preferences(tmp_db_path, [
            {
                "preference_text": "Already consumed preference",
                "preference_type": "phrasing",
                "applied_at": "2026-02-01T00:00:00",
            },
        ])

        conn = _sqlite3.connect(tmp_db_path)
        conn.row_factory = _sqlite3.Row

        captured_messages = []

        def capturing_call_claude(**kwargs):
            captured_messages.extend(kwargs.get("messages", []))
            return (
                {
                    "name": "Jane Doe",
                    "contact_line": "",
                    "summary": "DS.",
                    "skills": [],
                    "positions": [],
                    "education": [],
                },
                0.01,
            )

        with patch("job_finder.web.resume_generator.call_claude", side_effect=capturing_call_claude):
            with patch("job_finder.web.resume_generator.cost_gate", return_value=True):
                with patch("job_finder.web.resume_style_guide.load_style_guide", return_value={}):
                    mock_client = MagicMock()
                    generate_resume_single(
                        mock_client,
                        self._make_job_row(),
                        self._make_profile(),
                        conn,
                        self._make_config(),
                    )

        conn.close()

        assert captured_messages, "call_claude was not called"
        user_message = captured_messages[0]["content"]
        assert "Formatting Preferences" not in user_message, (
            "Should NOT include preferences section when all prefs have applied_at set"
        )
        assert "Already consumed preference" not in user_message, (
            "Consumed preference text should not appear in prompt"
        )

    def test_variant_path_includes_preferences_when_accepted_prefs_exist(self, tmp_db_path):
        """_generate_single_variant appends 'Formatting Preferences' when accepted prefs exist."""
        import sqlite3 as _sqlite3
        from unittest.mock import MagicMock, patch

        from job_finder.web.resume_generator import STRATEGY_POOL, _generate_single_variant

        self._make_db_with_preferences(tmp_db_path, [
            {"preference_text": "Use metrics in every bullet", "preference_type": "content_addition"},
        ])

        captured_messages = []

        def capturing_call_claude(**kwargs):
            captured_messages.extend(kwargs.get("messages", []))
            return (
                {
                    "name": "Jane Doe",
                    "contact_line": "",
                    "summary": "DS.",
                    "skills": [],
                    "positions": [],
                    "education": [],
                },
                0.01,
            )

        with patch("job_finder.web.resume_generator.call_claude", side_effect=capturing_call_claude):
            with patch("job_finder.web.resume_generator.cost_gate", return_value=True):
                mock_client = MagicMock()
                _generate_single_variant(
                    tmp_db_path,
                    lambda: mock_client,
                    self._make_job_row(),
                    self._make_profile(),
                    STRATEGY_POOL[0],
                    self._make_config(),
                )

        assert captured_messages, "call_claude was not called"
        user_message = captured_messages[0]["content"]
        assert "Formatting Preferences" in user_message, (
            "Expected '## Formatting Preferences' section in variant path user_message"
        )
        assert "Use metrics in every bullet" in user_message, (
            "Expected preference text in variant path user_message"
        )


# ---------------------------------------------------------------------------
# Style Guide injection tests (Phase 17-03)
# ---------------------------------------------------------------------------


class TestStyleGuideInjection:
    """Tests that style guide directives are injected into resume prompts."""

    def _make_job_row(self):
        from unittest.mock import MagicMock
        row = MagicMock()
        row.__getitem__ = lambda self, k: {
            "job_id": 1,
            "title": "Data Scientist",
            "company": "Acme Corp",
            "location": "Remote",
            "jd_full": "We need a data scientist with Python skills.",
            "platform": "linkedin",
            "url": "https://example.com/job/1",
        }.get(k, None)
        return row

    def _make_profile(self):
        return {
            "name": "Jane Doe",
            "contact_line": "jane@example.com",
            "summary": "Experienced data scientist.",
            "skills": ["Python", "SQL"],
            "positions": [],
            "education": [],
        }

    def _make_config(self):
        return {
            "scoring": {
                "models": {"sonnet": "claude-sonnet-4-6"},
                "monthly_budget_usd": 50.0,
            }
        }

    def _make_db(self, tmp_db_path):
        """Set up DB with required tables but no preferences."""
        import sqlite3 as _sqlite3
        from job_finder.web.db_migrate import run_migrations
        run_migrations(tmp_db_path)
        conn = _sqlite3.connect(tmp_db_path)
        conn.row_factory = _sqlite3.Row
        return conn

    def _canned_resume_response(self):
        return (
            {
                "name": "Jane Doe",
                "contact_line": "",
                "summary": "DS.",
                "skills": [],
                "positions": [],
                "education": [],
            },
            0.01,
        )

    def test_style_guide_directives_appear_in_generate_resume_single(self, tmp_db_path):
        """generate_resume_single includes style guide directives in Formatting Preferences."""
        import sqlite3 as _sqlite3
        from unittest.mock import MagicMock, patch

        from job_finder.web.resume_generator import generate_resume_single

        conn = self._make_db(tmp_db_path)
        captured_messages = []

        def capturing_call_claude(**kwargs):
            captured_messages.extend(kwargs.get("messages", []))
            return self._canned_resume_response()

        guide = {
            "bullet_style": "dashes",
            "verb_tense": "past",
            "section_order": ["Summary", "Experience", "Skills"],
            "tone": "direct",
            "date_format": "MMM YYYY",
        }

        with patch("job_finder.web.resume_generator.call_claude", side_effect=capturing_call_claude):
            with patch("job_finder.web.resume_generator.cost_gate", return_value=True):
                with patch("job_finder.web.resume_style_guide.load_style_guide", return_value=guide):
                    mock_client = MagicMock()
                    generate_resume_single(
                        mock_client,
                        self._make_job_row(),
                        self._make_profile(),
                        conn,
                        self._make_config(),
                    )

        conn.close()

        assert captured_messages, "call_claude was not called"
        user_message = captured_messages[0]["content"]
        assert "Formatting Preferences" in user_message, (
            "Expected '## Formatting Preferences' section from style guide"
        )
        assert "dashes" in user_message, "Expected bullet_style value in prompt"

    def test_style_guide_absent_no_formatting_section_in_generate_resume_single(self, tmp_db_path):
        """generate_resume_single omits Formatting Preferences when style guide is empty and no accepted prefs."""
        import sqlite3 as _sqlite3
        from unittest.mock import MagicMock, patch

        from job_finder.web.resume_generator import generate_resume_single

        conn = self._make_db(tmp_db_path)
        captured_messages = []

        def capturing_call_claude(**kwargs):
            captured_messages.extend(kwargs.get("messages", []))
            return self._canned_resume_response()

        with patch("job_finder.web.resume_generator.call_claude", side_effect=capturing_call_claude):
            with patch("job_finder.web.resume_generator.cost_gate", return_value=True):
                with patch("job_finder.web.resume_style_guide.load_style_guide", return_value={}):
                    mock_client = MagicMock()
                    generate_resume_single(
                        mock_client,
                        self._make_job_row(),
                        self._make_profile(),
                        conn,
                        self._make_config(),
                    )

        conn.close()

        assert captured_messages, "call_claude was not called"
        user_message = captured_messages[0]["content"]
        assert "Formatting Preferences" not in user_message, (
            "Should NOT include Formatting Preferences when style guide is empty and no prefs"
        )

    def test_style_guide_directives_appear_in_generate_single_variant(self, tmp_db_path):
        """_generate_single_variant includes style guide directives in Formatting Preferences."""
        from unittest.mock import MagicMock, patch

        from job_finder.web.resume_generator import STRATEGY_POOL, _generate_single_variant

        conn = self._make_db(tmp_db_path)
        conn.close()

        captured_messages = []

        def capturing_call_claude(**kwargs):
            captured_messages.extend(kwargs.get("messages", []))
            return self._canned_resume_response()

        guide = {
            "bullet_style": "bullets",
            "tone": "professional",
        }

        with patch("job_finder.web.resume_generator.call_claude", side_effect=capturing_call_claude):
            with patch("job_finder.web.resume_generator.cost_gate", return_value=True):
                with patch("job_finder.web.resume_style_guide.load_style_guide", return_value=guide):
                    mock_client = MagicMock()
                    _generate_single_variant(
                        tmp_db_path,
                        lambda: mock_client,
                        self._make_job_row(),
                        self._make_profile(),
                        STRATEGY_POOL[0],
                        self._make_config(),
                    )

        assert captured_messages, "call_claude was not called"
        user_message = captured_messages[0]["content"]
        assert "Formatting Preferences" in user_message, (
            "Expected '## Formatting Preferences' section in variant path from style guide"
        )
        assert "bullets" in user_message, "Expected bullet_style value in variant prompt"


class TestResumeGuidelines:
    """_RESUME_GUIDELINES constant exists and is integrated into _SYSTEM_PROMPT."""

    def test_guidelines_constant_exists_and_is_string(self):
        """_RESUME_GUIDELINES is importable from resume_generator and is a string."""
        from job_finder.web.resume_generator import _RESUME_GUIDELINES

        assert isinstance(_RESUME_GUIDELINES, str), (
            "_RESUME_GUIDELINES must be a string"
        )

    def test_guidelines_constant_min_length(self):
        """_RESUME_GUIDELINES is at least 500 characters (substantive content)."""
        from job_finder.web.resume_generator import _RESUME_GUIDELINES

        assert len(_RESUME_GUIDELINES) > 500, (
            f"_RESUME_GUIDELINES too short ({len(_RESUME_GUIDELINES)} chars); expected > 500"
        )

    def test_guidelines_contains_source_fidelity_language(self):
        """_RESUME_GUIDELINES contains source fidelity / fabrication prohibition language."""
        from job_finder.web.resume_generator import _RESUME_GUIDELINES

        lower = _RESUME_GUIDELINES.lower()
        assert "fabricat" in lower or "never list a skill" in lower, (
            "Expected source fidelity language ('fabricat' or 'never list a skill') in _RESUME_GUIDELINES"
        )

    def test_guidelines_contains_bullet_formula_language(self):
        """_RESUME_GUIDELINES contains bullet writing formula (action verb)."""
        from job_finder.web.resume_generator import _RESUME_GUIDELINES

        assert "action verb" in _RESUME_GUIDELINES.lower() or "Action Verb" in _RESUME_GUIDELINES, (
            "Expected bullet formula language ('action verb' or 'Action Verb') in _RESUME_GUIDELINES"
        )

    def test_guidelines_contains_seniority_count_language(self):
        """_RESUME_GUIDELINES contains seniority bullet count rules (4-6 bullets)."""
        from job_finder.web.resume_generator import _RESUME_GUIDELINES

        assert "4-6" in _RESUME_GUIDELINES, (
            "Expected seniority count language ('4-6') in _RESUME_GUIDELINES"
        )

    def test_guidelines_contains_typography_rules(self):
        """_RESUME_GUIDELINES contains typography rules (em dash prohibition)."""
        from job_finder.web.resume_generator import _RESUME_GUIDELINES

        lower = _RESUME_GUIDELINES.lower()
        assert "em dash" in lower, (
            "Expected typography rule ('em dash') in _RESUME_GUIDELINES"
        )

    def test_guidelines_contains_jd_mirroring_language(self):
        """_RESUME_GUIDELINES contains JD mirroring language."""
        from job_finder.web.resume_generator import _RESUME_GUIDELINES

        lower = _RESUME_GUIDELINES.lower()
        assert "mirror" in lower or "verbatim" in lower, (
            "Expected JD mirroring language ('mirror' or 'verbatim') in _RESUME_GUIDELINES"
        )

    def test_guidelines_contains_confidentiality_rules(self):
        """_RESUME_GUIDELINES contains confidentiality rules (client name)."""
        from job_finder.web.resume_generator import _RESUME_GUIDELINES

        lower = _RESUME_GUIDELINES.lower()
        assert "client name" in lower, (
            "Expected confidentiality rule ('client name') in _RESUME_GUIDELINES"
        )

    def test_guidelines_contains_skills_section_rules(self):
        """_RESUME_GUIDELINES contains skills section rules (no soft skills)."""
        from job_finder.web.resume_generator import _RESUME_GUIDELINES

        lower = _RESUME_GUIDELINES.lower()
        assert "soft skill" in lower or "hard skills only" in lower, (
            "Expected skills section rules ('soft skill' or 'hard skills only') in _RESUME_GUIDELINES"
        )

    def test_system_prompt_contains_guidelines(self):
        """_SYSTEM_PROMPT includes _RESUME_GUIDELINES content (guidelines injected into prompt)."""
        from job_finder.web.resume_generator import _RESUME_GUIDELINES, _SYSTEM_PROMPT

        assert _RESUME_GUIDELINES in _SYSTEM_PROMPT, (
            "_RESUME_GUIDELINES must be concatenated into _SYSTEM_PROMPT"
        )

    def test_system_prompt_retains_closed_world_constraint(self):
        """_SYSTEM_PROMPT still contains the original CRITICAL CONSTRAINT language."""
        from job_finder.web.resume_generator import _SYSTEM_PROMPT

        assert "CRITICAL CONSTRAINT" in _SYSTEM_PROMPT, (
            "Original 'CRITICAL CONSTRAINT' language must be retained in _SYSTEM_PROMPT"
        )


# =============================================================================
# Phase 33 Plan 01 Tests: Validation badge display in resume history
# =============================================================================


class TestValidationBadge:
    """Resume history entries show colored validation badge from validation_report JSON."""

    @pytest.fixture
    def badge_app(self, tmp_db_path):
        """Create test app with a Sonnet-scored job and resume_generations fixture data."""
        import sqlite3 as _sqlite3
        from job_finder.web.db_migrate import run_migrations
        from job_finder.web import create_app

        run_migrations(tmp_db_path)

        conn = _sqlite3.connect(tmp_db_path)
        conn.execute(
            """INSERT INTO jobs
                (dedup_key, title, company, location, sources, source_urls,
                 source_id, first_seen, last_seen, score, score_breakdown,
                 user_interest, pipeline_status, sonnet_score, jd_full)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
            (
                "badge|test|job",
                "Staff Data Scientist",
                "Badge Corp",
                "Remote",
                '["linkedin"]',
                '["https://linkedin.com/jobs/999"]',
                "999",
                "2026-03-01T00:00:00",
                "2026-03-17T00:00:00",
                9.0,
                "{}",
                "reviewing",
                "reviewing",
                88.0,
                "Full job description.",
            ),
        )
        conn.commit()
        conn.close()

        cfg = {
            "db": {"path": tmp_db_path},
            "scoring": {"models": {"sonnet": "claude-sonnet-4-6"}, "monthly_budget_usd": 25.0},
        }
        app = create_app(config=cfg)
        app.config["TESTING"] = True
        return app, tmp_db_path

    def _insert_resume_gen(self, db_path, validation_report=None, status="done",
                           doc_url="https://docs.google.com/doc/test"):
        """Insert a resume_generations row for badge|test|job."""
        import sqlite3 as _sqlite3
        conn = _sqlite3.connect(db_path)
        conn.execute(
            "INSERT INTO resume_generations "
            "(job_id, generated_at, model, status, doc_url, generation_type, validation_report) "
            "VALUES (?, ?, ?, ?, ?, ?, ?)",
            (
                "badge|test|job",
                "2026-03-17T00:00:00",
                "claude-sonnet-4-6",
                status,
                doc_url,
                "single",
                validation_report,
            ),
        )
        conn.commit()
        conn.close()

    def _get_expand_response(self, app, dedup_key="badge|test|job"):
        """GET the expand fragment for a job key."""
        from urllib.parse import quote
        from unittest.mock import patch
        with patch("job_finder.web.blueprints.jobs.get_drive_status", return_value={"ok": True}):
            with app.test_client() as client:
                resp = client.get(
                    f"/jobs/{quote(dedup_key, safe='')}/expand",
                    headers={"HX-Request": "true"},
                )
        return resp

    def test_badge_passed(self, badge_app):
        """Expand renders emerald-400 Passed badge when validation_report has no violations."""
        app, db_path = badge_app
        self._insert_resume_gen(
            db_path,
            validation_report='{"passed": true, "violations": [], "fix_summary": ""}',
        )

        resp = self._get_expand_response(app)

        assert resp.status_code == 200, f"Expected 200, got: {resp.status_code}"
        body = resp.data.decode()
        assert "text-emerald-400" in body, "Expected emerald-400 class for passed badge"
        assert "&#10003;" in body, "Expected checkmark entity for passed badge"
        assert "Passed" in body, "Expected 'Passed' text in badge"

    def test_badge_warnings(self, badge_app):
        """Expand renders amber-400 N warnings badge when validation_report has warnings only."""
        app, db_path = badge_app
        self._insert_resume_gen(
            db_path,
            validation_report='{"passed": true, "violations": [{"severity": "warning", "category": "style", "description": "test warning"}], "fix_summary": ""}',
        )

        resp = self._get_expand_response(app)

        assert resp.status_code == 200, f"Expected 200, got: {resp.status_code}"
        body = resp.data.decode()
        assert "text-amber-400" in body, "Expected amber-400 class for warnings badge"
        assert "&#9888;" in body, "Expected warning triangle entity for warnings badge"
        assert "1 warnings" in body, "Expected '1 warnings' count in badge"

    def test_badge_failed(self, badge_app):
        """Expand renders red-400 Failed badge when validation_report has errors."""
        app, db_path = badge_app
        self._insert_resume_gen(
            db_path,
            validation_report='{"passed": false, "violations": [{"severity": "error", "category": "content", "description": "test error"}], "fix_summary": ""}',
        )

        resp = self._get_expand_response(app)

        assert resp.status_code == 200, f"Expected 200, got: {resp.status_code}"
        body = resp.data.decode()
        assert "text-red-400" in body, "Expected red-400 class for failed badge"
        assert "&#10007;" in body, "Expected cross entity for failed badge"
        assert "Failed (1 errors)" in body, "Expected 'Failed (1 errors)' text in badge"

    def test_no_badge_when_null(self, badge_app):
        """Expand renders NO badge when validation_report is NULL."""
        app, db_path = badge_app
        self._insert_resume_gen(db_path, validation_report=None)

        resp = self._get_expand_response(app)

        assert resp.status_code == 200, f"Expected 200, got: {resp.status_code}"
        body = resp.data.decode()
        # Should NOT contain any badge color classes from validation
        assert "&#10003;" not in body, "Should not render checkmark when validation_report is null"
        assert "&#10007;" not in body, "Should not render cross when validation_report is null"
        assert "&#9888;" not in body, "Should not render warning triangle when validation_report is null"

    def test_violation_list_renders(self, badge_app):
        """Expand renders violation list with category:description format and correct color."""
        app, db_path = badge_app
        self._insert_resume_gen(
            db_path,
            validation_report='{"passed": true, "violations": [{"severity": "warning", "category": "style", "description": "test warning description"}], "fix_summary": ""}',
        )

        resp = self._get_expand_response(app)

        assert resp.status_code == 200, f"Expected 200, got: {resp.status_code}"
        body = resp.data.decode()
        assert "style: test warning description" in body, (
            "Expected 'category: description' format in violation list"
        )
        assert "text-amber-400" in body, "Expected amber color for warning violation"


# =============================================================================
# Phase 33 Plan 04 Tests: Validation badge in status polling done fragment
# =============================================================================


class TestStatusPollingBadge:
    """Status polling done fragment shows colored validation badge from validation_report JSON."""

    @pytest.fixture
    def polling_app(self, tmp_db_path):
        """Create test app with a Sonnet-scored job and resume_generations fixture data."""
        import sqlite3 as _sqlite3
        from job_finder.web.db_migrate import run_migrations
        from job_finder.web import create_app

        run_migrations(tmp_db_path)

        conn = _sqlite3.connect(tmp_db_path)
        conn.execute(
            """INSERT INTO jobs
                (dedup_key, title, company, location, sources, source_urls,
                 source_id, first_seen, last_seen, score, score_breakdown,
                 user_interest, pipeline_status, sonnet_score, jd_full)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
            (
                "polling|test|job",
                "Staff Data Scientist",
                "Polling Corp",
                "Remote",
                '["linkedin"]',
                '["https://linkedin.com/jobs/888"]',
                "888",
                "2026-03-01T00:00:00",
                "2026-03-17T00:00:00",
                9.0,
                "{}",
                "reviewing",
                "reviewing",
                88.0,
                "Full job description.",
            ),
        )
        conn.commit()
        conn.close()

        cfg = {
            "db": {"path": tmp_db_path},
            "scoring": {"models": {"sonnet": "claude-sonnet-4-6"}, "monthly_budget_usd": 25.0},
        }
        app = create_app(config=cfg)
        app.config["TESTING"] = True
        return app, tmp_db_path

    def _insert_resume_gen(self, db_path, validation_report=None, status="done",
                           doc_url="https://docs.google.com/doc/test"):
        """Insert a resume_generations row for polling|test|job and return gen_id."""
        import sqlite3 as _sqlite3
        conn = _sqlite3.connect(db_path)
        cursor = conn.execute(
            "INSERT INTO resume_generations "
            "(job_id, generated_at, model, status, doc_url, generation_type, validation_report) "
            "VALUES (?, ?, ?, ?, ?, ?, ?)",
            (
                "polling|test|job",
                "2026-03-17T00:00:00",
                "claude-sonnet-4-6",
                status,
                doc_url,
                "single",
                validation_report,
            ),
        )
        gen_id = cursor.lastrowid
        conn.commit()
        conn.close()
        return gen_id

    def _get_status_response(self, app, db_path, validation_report=None):
        """Insert done resume_gen and GET the status polling fragment."""
        from urllib.parse import quote
        gen_id = self._insert_resume_gen(db_path, validation_report=validation_report)
        dedup_key = "polling|test|job"
        with app.test_client() as client:
            resp = client.get(f"/jobs/{quote(dedup_key, safe='')}/resume/status/{gen_id}")
        return resp

    def test_status_done_shows_passed_badge(self, polling_app):
        """Status polling done fragment shows emerald-400 Passed badge when validation passes."""
        app, db_path = polling_app
        resp = self._get_status_response(
            app, db_path,
            validation_report='{"passed": true, "violations": [], "fix_summary": ""}',
        )
        assert resp.status_code == 200, f"Expected 200, got: {resp.status_code}"
        body = resp.data.decode()
        assert "text-emerald-400" in body, "Expected emerald-400 class for passed badge"
        assert "&#10003;" in body, "Expected checkmark entity for passed badge"
        assert "Passed" in body, "Expected 'Passed' text in badge"

    def test_status_done_shows_warning_badge(self, polling_app):
        """Status polling done fragment shows amber-400 warnings badge when warnings present."""
        app, db_path = polling_app
        resp = self._get_status_response(
            app, db_path,
            validation_report='{"passed": true, "violations": [{"severity": "warning", "category": "style", "description": "test warning"}], "fix_summary": ""}',
        )
        assert resp.status_code == 200, f"Expected 200, got: {resp.status_code}"
        body = resp.data.decode()
        assert "text-amber-400" in body, "Expected amber-400 class for warnings badge"
        assert "&#9888;" in body, "Expected warning triangle entity for warnings badge"

    def test_status_done_shows_failed_badge(self, polling_app):
        """Status polling done fragment shows red-400 Failed badge when errors present."""
        app, db_path = polling_app
        resp = self._get_status_response(
            app, db_path,
            validation_report='{"passed": false, "violations": [{"severity": "error", "category": "content", "description": "test error"}], "fix_summary": ""}',
        )
        assert resp.status_code == 200, f"Expected 200, got: {resp.status_code}"
        body = resp.data.decode()
        assert "text-red-400" in body, "Expected red-400 class for failed badge"
        assert "&#10007;" in body, "Expected cross entity for failed badge"

    def test_status_done_no_badge_when_null(self, polling_app):
        """Status polling done fragment shows no badge when validation_report is NULL."""
        app, db_path = polling_app
        resp = self._get_status_response(app, db_path, validation_report=None)
        assert resp.status_code == 200, f"Expected 200, got: {resp.status_code}"
        body = resp.data.decode()
        assert "&#10003;" not in body, "Should not render checkmark when validation_report is null"
        assert "&#10007;" not in body, "Should not render cross when validation_report is null"
        assert "&#9888;" not in body, "Should not render warning triangle when validation_report is null"
